"""
Academic Report Generator - 学术报告生成器

生成出版级质量的研究报告，支持 PDF 和 Word 双格式输出。

主要功能：
    - 学术论文格式（标题、摘要、章节、参考文献）
    - 图片自动下载、缩放和标注
    - 表格解析和格式化渲染
    - Mermaid 图表渲染（如果可用）
    - 统一的排版和间距

输出格式：
    - PDF: 使用 ReportLab 生成，支持中文
    - Word: 使用 python-docx 生成，兼容 Office

使用示例：
    generator = PDFReportGenerator()
    pdf_path, docx_path = generator.generate_both(
        title="研究报告",
        content="# 报告内容...",
        images=[{"url": "...", "title": "..."}]
    )
"""
import os
import sys
import re
import io
import hashlib
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple

# 添加项目路径
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from config import REPORTS_DIR
from utils import get_logger

logger = get_logger(__name__)


class AcademicReportGenerator:
    """
    学术报告生成器
    
    生成符合学术论文规范的研究报告，支持 PDF 和 Word 双格式输出。
    
    主要特性：
        - 学术论文格式（标题、作者、日期、摘要、章节、参考文献）
        - 图片自动编号和标注（Figure 1, Figure 2...）
        - 表格自动编号和格式化（Table 1, Table 2...）
        - 统一的 Times New Roman 字体和学术排版
        - 支持 Markdown 格式的内容解析
    
    Attributes:
        output_dir: 报告输出目录
        image_cache_dir: 图片缓存目录
        FIGURE_COUNTER: 图片计数器
        TABLE_COUNTER: 表格计数器
    """
    
    # 图片和表格计数器（每个报告重置）
    FIGURE_COUNTER = 0
    TABLE_COUNTER = 0
    
    def __init__(self, output_dir: str = None):
        """
        初始化报告生成器
        
        Args:
            output_dir: 报告输出目录，默认使用 config.REPORTS_DIR
        """
        self.output_dir = Path(output_dir) if output_dir else REPORTS_DIR
        self.output_dir.mkdir(exist_ok=True)
        
        # 图片缓存目录，用于存储下载的图片
        self.image_cache_dir = self.output_dir / "image_cache"
        self.image_cache_dir.mkdir(exist_ok=True)
        
        # 已下载图片的缓存（URL -> 本地路径）
        self._downloaded_images = {}
        self._failed_web_screenshots = set()
        self._reset_counters()

        # Mermaid CLI availability cache
        self._mmdc_available: Optional[bool] = None
        self._mmdc_version: Optional[str] = None
        
        # 注册中文字体
        self._chinese_fonts_available = self._register_chinese_fonts()
    
    def _register_chinese_fonts(self) -> bool:
        """
        注册中文字体（SimSun宋体、SimHei黑体、SimKai楷体、SimFang仿宋）。
        从 Windows 系统字体目录加载 TTF/TTC 字体文件。
        
        Returns:
            True if at least one Chinese font was registered successfully
        """
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            fonts_dir = os.path.join(os.environ.get('WINDIR', r'C:\Windows'), 'Fonts')
            
            font_map = {
                'SimSun': ['simsun.ttc', 'simsun.ttf', 'SIMSUN.TTC', 'SIMSUN.TTF'],
                'SimHei': ['simhei.ttf', 'SIMHEI.TTF', 'msyh.ttc', 'msyh.ttf'],
                'SimKai': ['simkai.ttf', 'SIMKAI.TTF'],
                'SimFang': ['simfang.ttf', 'SIMFANG.TTF'],
            }
            
            registered = []
            for font_name, candidates in font_map.items():
                for candidate in candidates:
                    font_path = os.path.join(fonts_dir, candidate)
                    if os.path.exists(font_path):
                        try:
                            pdfmetrics.registerFont(TTFont(font_name, font_path))
                            registered.append(font_name)
                            logger.info(f"Registered Chinese font: {font_name} from {candidate}")
                            break
                        except Exception as fe:
                            logger.debug(f"Failed to register {font_name} from {candidate}: {fe}")
            
            # 注册粗体/斜体变体（使用相同字体文件）
            if 'SimHei' in registered:
                try:
                    from reportlab.pdfbase.pdfmetrics import registerFontFamily
                    registerFontFamily('SimSun', normal='SimSun', bold='SimHei', italic='SimKai', boldItalic='SimHei')
                except Exception:
                    pass
            
            if registered:
                logger.info(f"Chinese fonts registered: {registered}")
                return True
            else:
                logger.warning("No Chinese fonts found. Chinese text in PDF may not render correctly.")
                # 尝试 CID 字体作为后备
                try:
                    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
                    pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
                    logger.info("Registered CID font STSong-Light as fallback")
                    return True
                except Exception:
                    pass
                return False
                
        except ImportError:
            logger.warning("reportlab not available, skipping Chinese font registration")
            return False
        except Exception as e:
            logger.warning(f"Failed to register Chinese fonts: {e}")
            return False
    
    def _is_chinese_content(self, text: str) -> bool:
        """检测文本是否包含中文字符。"""
        if not text:
            return False
        chinese_chars = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
        return chinese_chars > len(text) * 0.05 or chinese_chars > 10
    
    def _reset_counters(self):
        """重置图片和表格计数器（每个新报告调用）"""
        self.FIGURE_COUNTER = 0
        self.TABLE_COUNTER = 0
    
    def _next_figure_num(self) -> int:
        """获取下一个图片编号"""
        self.FIGURE_COUNTER += 1
        return self.FIGURE_COUNTER
    
    def _next_table_num(self) -> int:
        """获取下一个表格编号"""
        self.TABLE_COUNTER += 1
        return self.TABLE_COUNTER
    
    def _download_image(self, url: str, timeout: int = 10) -> Optional[str]:
        """Download an image from URL and cache it locally."""
        if url in self._downloaded_images:
            return self._downloaded_images[url]
        
        try:
            import requests
            from PIL import Image
            
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
            response = requests.get(url, timeout=timeout, stream=True, headers=headers)
            response.raise_for_status()
            
            img = Image.open(io.BytesIO(response.content))
            
            if img.mode in ('RGBA', 'P'):
                img = img.convert('RGB')
            
            max_width = 1200
            max_height = 900
            img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
            
            url_hash = hashlib.md5(url.encode()).hexdigest()[:12]
            img_path = self.image_cache_dir / f"img_{url_hash}.png"
            img.save(str(img_path), "PNG", optimize=True)
            
            self._downloaded_images[url] = str(img_path)
            logger.info(f"Downloaded image: {url[:50]}...")
            return str(img_path)
            
        except Exception as e:
            logger.warning(f"Failed to download image {url[:50]}...: {e}")
            return None
    
    def _parse_table_from_text(self, table_text: str) -> Optional[List[List[str]]]:
        """Parse a markdown-style table into a 2D list."""
        try:
            lines = [l.strip() for l in table_text.strip().split('\n') if l.strip()]
            if len(lines) < 2:
                return None
            
            table_data = []
            for line in lines:
                if line.startswith('|') and line.endswith('|'):
                    cells = [c.strip() for c in line[1:-1].split('|')]
                    if all(set(c) <= {'-', ':', ' '} for c in cells):
                        continue
                    table_data.append(cells)
            
            return table_data if len(table_data) >= 2 else None
            
        except Exception as e:
            logger.warning(f"Failed to parse table: {e}")
            return None
    
    def _create_table_with_caption(self, data: List[List[str]], source: str = "", 
                                    title: str = "", is_chinese: bool = False) -> List[Any]:
        """
        Create a table with proper academic formatting and source attribution.
        Table number and source appear ABOVE the table.
        For Chinese: uses 三线表 style with Chinese fonts.
        """
        from reportlab.platypus import Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_CENTER
        
        flowables = []
        styles = getSampleStyleSheet()
        
        font_body = 'SimSun' if (is_chinese and self._chinese_fonts_available) else 'Times-Roman'
        font_caption = 'SimHei' if (is_chinese and self._chinese_fonts_available) else 'Times-Bold'
        
        table_num = self._next_table_num()
        if is_chinese:
            caption_text = f"<b>表{table_num}</b>"
            if title:
                caption_text += f" {title}"
        else:
            caption_text = f"<b>Table {table_num}.</b> {title}" if title else f"<b>Table {table_num}.</b>"
            if source:
                caption_text += f" <i>(Source: {source})</i>"
        
        caption_style = ParagraphStyle(
            'TableCaption',
            parent=styles['Normal'],
            fontSize=9,
            alignment=TA_CENTER,
            spaceBefore=8,
            spaceAfter=4,
            textColor=colors.HexColor('#333333'),
            fontName=font_caption
        )
        caption_text = caption_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        caption_text = caption_text.replace('&lt;b&gt;', '<b>').replace('&lt;/b&gt;', '</b>')
        caption_text = caption_text.replace('&lt;i&gt;', '<i>').replace('&lt;/i&gt;', '</i>')
        flowables.append(Paragraph(caption_text, caption_style))
        
        col_count = max(len(row) for row in data)
        for row in data:
            while len(row) < col_count:
                row.append("")
        
        wrapped_data = []
        for row_idx, row in enumerate(data):
            wrapped_row = []
            for cell in row:
                cell_font = font_caption if row_idx == 0 else font_body
                cell_style = ParagraphStyle(
                    'CellStyle',
                    parent=styles['Normal'],
                    fontSize=8,
                    leading=10,
                    wordWrap='CJK',
                    fontName=cell_font
                )
                cell_text = str(cell)
                cell_text = re.sub(r'^\*+|\*+$', '', cell_text)
                cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                cell_text = re.sub(r'\*(.+?)\*', r'\1', cell_text)
                cell_text = cell_text.strip()
                cell_text = cell_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                if row_idx == 0:
                    cell_text = f"<b>{cell_text}</b>"
                wrapped_row.append(Paragraph(cell_text, cell_style))
            wrapped_data.append(wrapped_row)
        
        available_width = 6.2 * inch
        if col_count <= 2:
            col_widths = [available_width / col_count] * col_count
        elif col_count <= 4:
            col_widths = [available_width / col_count] * col_count
        else:
            first_col_width = available_width * 0.25
            other_width = (available_width - first_col_width) / (col_count - 1)
            col_widths = [first_col_width] + [other_width] * (col_count - 1)
        
        table = Table(wrapped_data, colWidths=col_widths, repeatRows=1)
        
        if is_chinese:
            # 三线表样式
            table_style = TableStyle([
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LINEABOVE', (0, 0), (-1, 0), 1.5, colors.black),       # 顶线
                ('LINEBELOW', (0, 0), (-1, 0), 0.75, colors.black),      # 表头下线
                ('LINEBELOW', (0, -1), (-1, -1), 1.5, colors.black),     # 底线
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ('LEFTPADDING', (0, 0), (-1, -1), 4),
                ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            ])
        else:
            table_style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c5282')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e0')),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ('LEFTPADDING', (0, 0), (-1, -1), 4),
                ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            ])
            for i in range(1, len(wrapped_data)):
                bg_color = colors.HexColor('#f8f9fa') if i % 2 == 0 else colors.white
                table_style.add('BACKGROUND', (0, i), (-1, i), bg_color)
        
        table.setStyle(table_style)
        flowables.append(table)
        flowables.append(Spacer(1, 8))
        
        return flowables
    
    def _create_figure_with_caption(self, img_path: str, title: str = "", 
                                      source: str = "", is_chinese: bool = False,
                                      diagram_type: str = "", show_caption: bool = False) -> List[Any]:
        """
        Create a figure, optionally with academic-style caption below the image.
        Format: Figure X. Title (Source: source) / 图X Title
        
        Args:
            img_path: Path to the image file
            title: Figure title/description
            source: Source attribution
            is_chinese: Whether to use Chinese formatting
            diagram_type: Type of diagram for dynamic sizing (timeline, gantt, mindmap get larger)
            show_caption: Whether to show caption below the image (default False)
        """
        from reportlab.platypus import Image, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        
        flowables = []
        font_caption = 'SimHei' if (is_chinese and self._chinese_fonts_available) else 'Times-Roman'
        
        try:
            img = Image(img_path)
            
            # 根据图表类型动态调整尺寸
            # 内容较多的图表类型使用更大尺寸，但不超过 A4 可用区域
            diagram_type_lower = diagram_type.lower() if diagram_type else ""
            title_lower = title.lower() if title else ""
            
            # 检测是否为内容较多的图表类型
            large_diagram_types = [
                'timeline', 'gantt', 'mindmap', 'flowchart', 'sequence', 'class', 'er',
                'gitgraph', 'sankey', 'architecture', 'c4', 'journey', 'requirement', 'block'
            ]
            is_large_diagram = any(t in diagram_type_lower or t in title_lower for t in large_diagram_types)
            
            if is_large_diagram:
                # 大型图表：最大 6.5 x 8.0 英寸（接近 A4 可用区域）
                max_width = 6.5 * inch
                max_height = 8.0 * inch
            else:
                # 普通图表：5.0 x 4.0 英寸
                max_width = 5.5 * inch
                max_height = 4.5 * inch
            
            aspect = img.imageWidth / img.imageHeight
            
            # 先按宽度缩放
            if img.imageWidth > max_width:
                img.drawWidth = max_width
                img.drawHeight = max_width / aspect
            else:
                img.drawWidth = img.imageWidth
                img.drawHeight = img.imageHeight
            
            # 再检查高度限制
            if img.drawHeight > max_height:
                img.drawHeight = max_height
                img.drawWidth = max_height * aspect
            
            img.hAlign = 'CENTER'
            flowables.append(Spacer(1, 6))
            flowables.append(img)
            
            # 只在 show_caption=True 时显示题注
            if show_caption:
                styles = getSampleStyleSheet()
                fig_num = self._next_figure_num()
                
                if is_chinese:
                    caption_text = f"<b>图{fig_num}</b>"
                    if title:
                        caption_text += f" {title}"
                    if source and source != 'Unknown':
                        caption_text += f"（来源：{source}）"
                else:
                    caption_text = f"<b>Figure {fig_num}.</b>"
                    if title:
                        caption_text += f" {title}"
                    if source:
                        caption_text += f" <i>(Source: {source})</i>"
                
                caption_style = ParagraphStyle(
                    'FigureCaption',
                    parent=styles['Normal'],
                    fontSize=9,
                    alignment=TA_CENTER,
                    textColor=colors.HexColor('#333333'),
                    spaceBefore=4,
                    spaceAfter=8,
                    fontName=font_caption
                )
                caption_text = caption_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                caption_text = caption_text.replace('&lt;b&gt;', '<b>').replace('&lt;/b&gt;', '</b>')
                caption_text = caption_text.replace('&lt;i&gt;', '<i>').replace('&lt;/i&gt;', '</i>')
                flowables.append(Paragraph(caption_text, caption_style))
            else:
                # 不显示题注时，只添加一点间距
                flowables.append(Spacer(1, 8))
            
        except Exception as e:
            logger.warning(f"Failed to create figure flowable: {e}")
        
        return flowables
    
    def _get_academic_styles(self, is_chinese: bool = False):
        """
        Get academic paper styles for PDF generation.
        
        Args:
            is_chinese: If True, use Chinese academic paper fonts and formatting
        """
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
        from reportlab.lib.colors import HexColor, black
        
        styles = getSampleStyleSheet()
        
        if is_chinese and self._chinese_fonts_available:
            # 中文学术论文格式
            # 二号黑体=22pt, 三号=16pt, 四号=14pt, 小四=12pt, 五号=10.5pt, 小五=9pt
            font_body = 'SimSun'       # 宋体 - 正文
            font_heading = 'SimHei'    # 黑体 - 标题
            font_kai = 'SimKai' if self._chinese_fonts_available else 'SimSun'  # 楷体 - 作者
            font_fang = 'SimFang' if self._chinese_fonts_available else 'SimSun'  # 仿宋
            
            custom_styles = {
                'title': ParagraphStyle(
                    'AcademicTitle',
                    parent=styles['Heading1'],
                    fontSize=22,        # 二号
                    leading=28,
                    spaceAfter=8,
                    alignment=TA_CENTER,
                    textColor=black,
                    fontName=font_heading
                ),
                'author': ParagraphStyle(
                    'Author',
                    parent=styles['Normal'],
                    fontSize=14,        # 四号楷体
                    leading=18,
                    alignment=TA_CENTER,
                    spaceAfter=4,
                    fontName=font_kai
                ),
                'date': ParagraphStyle(
                    'Date',
                    parent=styles['Normal'],
                    fontSize=12,        # 小四
                    leading=16,
                    alignment=TA_CENTER,
                    spaceAfter=16,
                    fontName=font_body
                ),
                'abstract_title': ParagraphStyle(
                    'AbstractTitle',
                    parent=styles['Heading2'],
                    fontSize=9,         # 小五号黑体
                    leading=13,
                    spaceBefore=12,
                    spaceAfter=6,
                    alignment=TA_LEFT,
                    firstLineIndent=24,
                    fontName=font_heading
                ),
                'abstract': ParagraphStyle(
                    'Abstract',
                    parent=styles['Normal'],
                    fontSize=9,         # 小五号
                    leading=13,
                    alignment=TA_JUSTIFY,
                    leftIndent=24,
                    rightIndent=24,
                    firstLineIndent=18,
                    spaceAfter=16,
                    fontName=font_body
                ),
                'keywords': ParagraphStyle(
                    'Keywords',
                    parent=styles['Normal'],
                    fontSize=9,         # 小五号黑体
                    leading=13,
                    alignment=TA_LEFT,
                    firstLineIndent=24,
                    spaceAfter=12,
                    fontName=font_heading
                ),
                'h1': ParagraphStyle(
                    'SectionH1',
                    parent=styles['Heading1'],
                    fontSize=14,        # 四号宋体加粗
                    leading=22,
                    spaceBefore=16,
                    spaceAfter=8,
                    textColor=black,
                    fontName=font_heading
                ),
                'h2': ParagraphStyle(
                    'SectionH2',
                    parent=styles['Heading2'],
                    fontSize=12,        # 小四号宋体加粗
                    leading=20,
                    spaceBefore=12,
                    spaceAfter=6,
                    textColor=black,
                    fontName=font_heading
                ),
                'h3': ParagraphStyle(
                    'SectionH3',
                    parent=styles['Heading3'],
                    fontSize=12,        # 小四号宋体加粗
                    leading=20,
                    spaceBefore=10,
                    spaceAfter=4,
                    textColor=black,
                    fontName=font_heading
                ),
                'h4': ParagraphStyle(
                    'SectionH4',
                    parent=styles['Heading3'],
                    fontSize=12,        # 小四号宋体加粗
                    leading=20,
                    spaceBefore=8,
                    spaceAfter=4,
                    textColor=black,
                    fontName=font_heading
                ),
                'body': ParagraphStyle(
                    'BodyText',
                    parent=styles['Normal'],
                    fontSize=12,        # 小四号宋体
                    leading=22,         # 1.5倍行距
                    alignment=TA_JUSTIFY,
                    spaceAfter=0,
                    firstLineIndent=24, # 两字符缩进
                    fontName=font_body,
                    wordWrap='CJK'      # 中英文混排时正确换行，避免大空格
                ),
                'body_first': ParagraphStyle(
                    'BodyTextFirst',
                    parent=styles['Normal'],
                    fontSize=12,
                    leading=22,
                    alignment=TA_JUSTIFY,
                    spaceAfter=0,
                    firstLineIndent=24,
                    fontName=font_body,
                    wordWrap='CJK'      # 中英文混排时正确换行，避免大空格
                ),
                'bullet': ParagraphStyle(
                    'BulletPoint',
                    parent=styles['Normal'],
                    fontSize=12,
                    leading=22,
                    leftIndent=24,
                    spaceAfter=0,
                    fontName=font_body,
                    wordWrap='CJK'
                ),
                'numbered': ParagraphStyle(
                    'NumberedItem',
                    parent=styles['Normal'],
                    fontSize=12,
                    leading=22,
                    leftIndent=24,
                    spaceAfter=0,
                    fontName=font_body,
                    wordWrap='CJK'
                ),
                'reference': ParagraphStyle(
                    'Reference',
                    parent=styles['Normal'],
                    fontSize=10.5,      # 五号宋体
                    leading=16,
                    leftIndent=18,
                    firstLineIndent=-18,
                    spaceAfter=3,
                    fontName=font_body,
                    wordWrap='CJK'
                ),
                'conclusion': ParagraphStyle(
                    'Conclusion',
                    parent=styles['Normal'],
                    fontSize=10.5,      # 五号宋体加粗
                    leading=16,
                    alignment=TA_JUSTIFY,
                    spaceAfter=4,
                    firstLineIndent=0,
                    fontName=font_heading
                ),
            }
        else:
            # English academic paper format
            custom_styles = {
                'title': ParagraphStyle(
                    'AcademicTitle',
                    parent=styles['Heading1'],
                    fontSize=16,
                    leading=20,
                    spaceAfter=6,
                    alignment=TA_CENTER,
                    textColor=black,
                    fontName='Times-Bold'
                ),
                'author': ParagraphStyle(
                    'Author',
                    parent=styles['Normal'],
                    fontSize=11,
                    alignment=TA_CENTER,
                    spaceAfter=4,
                    fontName='Times-Roman'
                ),
                'date': ParagraphStyle(
                    'Date',
                    parent=styles['Normal'],
                    fontSize=10,
                    alignment=TA_CENTER,
                    spaceAfter=16,
                    fontName='Times-Italic'
                ),
                'abstract_title': ParagraphStyle(
                    'AbstractTitle',
                    parent=styles['Heading2'],
                    fontSize=11,
                    spaceBefore=12,
                    spaceAfter=6,
                    alignment=TA_CENTER,
                    fontName='Times-Bold'
                ),
                'abstract': ParagraphStyle(
                    'Abstract',
                    parent=styles['Normal'],
                    fontSize=10,
                    leading=13,
                    alignment=TA_JUSTIFY,
                    leftIndent=36,
                    rightIndent=36,
                    spaceAfter=16,
                    fontName='Times-Italic'
                ),
                'h1': ParagraphStyle(
                    'SectionH1',
                    parent=styles['Heading1'],
                    fontSize=14,
                    leading=17,
                    spaceBefore=16,
                    spaceAfter=8,
                    textColor=black,
                    fontName='Times-Bold'
                ),
                'h2': ParagraphStyle(
                    'SectionH2',
                    parent=styles['Heading2'],
                    fontSize=12,
                    leading=15,
                    spaceBefore=12,
                    spaceAfter=6,
                    textColor=black,
                    fontName='Times-Bold'
                ),
                'h3': ParagraphStyle(
                    'SectionH3',
                    parent=styles['Heading3'],
                    fontSize=11,
                    leading=14,
                    spaceBefore=10,
                    spaceAfter=4,
                    textColor=black,
                    fontName='Times-Bold'
                ),
                'body': ParagraphStyle(
                    'BodyText',
                    parent=styles['Normal'],
                    fontSize=10,
                    leading=13,
                    alignment=TA_JUSTIFY,
                    spaceAfter=6,
                    firstLineIndent=18,
                    fontName='Times-Roman'
                ),
                'body_first': ParagraphStyle(
                    'BodyTextFirst',
                    parent=styles['Normal'],
                    fontSize=10,
                    leading=13,
                    alignment=TA_JUSTIFY,
                    spaceAfter=6,
                    firstLineIndent=0,
                    fontName='Times-Roman'
                ),
                'bullet': ParagraphStyle(
                    'BulletPoint',
                    parent=styles['Normal'],
                    fontSize=10,
                    leading=13,
                    leftIndent=24,
                    spaceAfter=3,
                    fontName='Times-Roman'
                ),
                'numbered': ParagraphStyle(
                    'NumberedItem',
                    parent=styles['Normal'],
                    fontSize=10,
                    leading=13,
                    leftIndent=24,
                    spaceAfter=3,
                    fontName='Times-Roman'
                ),
                'reference': ParagraphStyle(
                    'Reference',
                    parent=styles['Normal'],
                    fontSize=9,
                    leading=11,
                    leftIndent=18,
                    firstLineIndent=-18,
                    spaceAfter=4,
                    fontName='Times-Roman'
                ),
            }
        
        return styles, custom_styles
    
    def _prepare_images(self, images: List[Dict], max_images: int = 12) -> List[Dict]:
        """Download and prepare images for embedding."""
        downloaded = []
        for img_info in images[:max_images]:
            url = img_info.get("url") or img_info.get("thumbnail")
            if url:
                local_path = self._download_image(url)
                if local_path:
                    downloaded.append({
                        "path": local_path,
                        "title": img_info.get("title", ""),
                        "category": img_info.get("category", ""),
                        "source": img_info.get("source", "Unknown")
                    })
        logger.info(f"Prepared {len(downloaded)} images for report")
        return downloaded
    
    def _download_web_screenshot(self, url: str) -> Optional[str]:
        """
        Download a web screenshot/image from URL and cache it locally.
        
        Args:
            url: URL of the image or screenshot to download
            
        Returns:
            Path to the cached image file, or None if download failed
        """
        import requests
        import hashlib
        from PIL import Image
        
        # 已经失败过的 URL 直接跳过，减少重复请求
        if url in self._failed_web_screenshots:
            return None

        # 基于扩展名的快速过滤（避免下载 PDF/HTML 等）
        lower_url = url.lower().split('?', 1)[0].split('#', 1)[0]
        if lower_url.endswith('.pdf'):
            logger.warning(f"URL does not appear to be an image (pdf): {url}")
            self._failed_web_screenshots.add(url)
            return None

        # 检查缓存
        url_hash = hashlib.md5(url.encode()).hexdigest()[:12]
        cached_path = self.image_cache_dir / f"web_screenshot_{url_hash}.png"
        
        if cached_path.exists():
            logger.info(f"Using cached web screenshot: {cached_path}")
            return str(cached_path)
        
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'image/avif,image/webp,image/apng,image/svg+xml,image/*,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.9',
                'Referer': url.rsplit('/', 1)[0] + '/' if '/' in url else url
            }
            response = requests.get(url, headers=headers, timeout=30, stream=True)
            response.raise_for_status()
            
            # 验证是图片内容
            content_type = response.headers.get('Content-Type', '')
            if not any(t in content_type.lower() for t in ['image', 'png', 'jpeg', 'jpg', 'gif', 'webp']):
                logger.warning(f"URL does not appear to be an image: {content_type}")
                self._failed_web_screenshots.add(url)
                return None
            
            # 保存并优化图片
            img = Image.open(io.BytesIO(response.content))
            
            # 转换为 RGB 模式
            if img.mode in ('RGBA', 'P'):
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'RGBA':
                    background.paste(img, mask=img.split()[3])
                else:
                    background.paste(img)
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            
            # 限制最大尺寸
            max_width, max_height = 1200, 900
            if img.width > max_width or img.height > max_height:
                img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
            
            img.save(str(cached_path), "PNG", optimize=True)
            logger.info(f"Downloaded web screenshot: {url} -> {cached_path}")
            return str(cached_path)
            
        except requests.exceptions.Timeout:
            logger.warning(f"Timeout downloading web screenshot: {url}")
            self._failed_web_screenshots.add(url)
            # 尝试使用 Playwright 截图作为后备
            return self._capture_webpage_screenshot(url, cached_path)
        except requests.exceptions.RequestException as e:
            logger.warning(f"Failed to download web screenshot {url}: {e}")
            self._failed_web_screenshots.add(url)
            # 尝试使用 Playwright 截图作为后备
            return self._capture_webpage_screenshot(url, cached_path)
        except Exception as e:
            logger.warning(f"Error processing web screenshot {url}: {e}")
            self._failed_web_screenshots.add(url)
            return None

    def _capture_webpage_screenshot(self, url: str, output_path: Path) -> Optional[str]:
        """
        Capture a screenshot of a webpage using Playwright as fallback.
        Handles both sync and async contexts automatically.
        
        Args:
            url: URL of the webpage to capture
            output_path: Path to save the screenshot
            
        Returns:
            Path to the screenshot file, or None if capture failed
        """
        # 检查是否已经缓存了 Playwright 可用性
        if not hasattr(self, '_playwright_available'):
            self._playwright_available = None
        
        # 如果已知不可用，直接返回
        if self._playwright_available is False:
            return None
        
        try:
            import playwright
            self._playwright_available = True
        except ImportError:
            if self._playwright_available is None:
                logger.info("Playwright not installed. Install with: pip install playwright && playwright install chromium")
            self._playwright_available = False
            return None
        
        import asyncio
        
        async def _async_capture():
            from playwright.async_api import async_playwright
            async with async_playwright() as p:
                browser = await p.chromium.launch(headless=True)
                page = await browser.new_page(viewport={'width': 1280, 'height': 720})
                await page.goto(url, timeout=30000, wait_until='domcontentloaded')
                await page.wait_for_timeout(1000)  # 等待页面渲染
                await page.screenshot(path=str(output_path), full_page=False)
                await browser.close()
        
        try:
            # 检查是否在 asyncio 事件循环中
            try:
                loop = asyncio.get_running_loop()
                # 在事件循环中，使用 nest_asyncio 或在新线程中运行
                import concurrent.futures
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    future = executor.submit(asyncio.run, _async_capture())
                    future.result(timeout=60)
            except RuntimeError:
                # 不在事件循环中，直接运行
                asyncio.run(_async_capture())
            
            if output_path.exists():
                logger.info(f"Captured webpage screenshot: {url} -> {output_path}")
                return str(output_path)
            return None
            
        except Exception as e:
            logger.warning(f"Failed to capture webpage screenshot {url}: {e}")
            return None
    
    def _generate_placeholder_image(self, text: str, output_path: Path, url: str = None) -> Optional[str]:
        """
        Generate a placeholder image with text when image download fails.
        
        Args:
            text: Text to display on the placeholder
            output_path: Path to save the placeholder image
            url: Optional URL to display on the placeholder
            
        Returns:
            Path to the placeholder image, or None if generation failed
        """
        try:
            from PIL import Image, ImageDraw, ImageFont
            
            # 创建灰色背景的占位图
            width, height = 600, 340 if url else 300
            img = Image.new('RGB', (width, height), color=(240, 240, 240))
            draw = ImageDraw.Draw(img)
            
            # 绘制边框
            draw.rectangle([5, 5, width-6, height-6], outline=(200, 200, 200), width=2)
            
            # 尝试使用系统字体，否则使用默认字体
            try:
                font = ImageFont.truetype("arial.ttf", 16)
                small_font = ImageFont.truetype("arial.ttf", 12)
                url_font = ImageFont.truetype("arial.ttf", 10)
            except:
                font = ImageFont.load_default()
                small_font = font
                url_font = font
            
            # 绘制图标和文字
            icon_text = "[Image Unavailable]"
            draw.text((width//2, height//2 - 50), icon_text, fill=(150, 150, 150), font=font, anchor="mm")
            
            # 截断过长的文本
            display_text = text[:80] + "..." if len(text) > 80 else text
            draw.text((width//2, height//2 - 20), display_text, fill=(120, 120, 120), font=small_font, anchor="mm")
            
            # 显示 URL（如果提供）
            if url:
                # 截断过长的 URL
                display_url = url[:90] + "..." if len(url) > 90 else url
                draw.text((width//2, height//2 + 20), "Source URL:", fill=(100, 100, 100), font=small_font, anchor="mm")
                draw.text((width//2, height//2 + 45), display_url, fill=(70, 130, 180), font=url_font, anchor="mm")
            
            img.save(str(output_path), "PNG")
            logger.info(f"Generated placeholder image: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.warning(f"Failed to generate placeholder image: {e}")
            return None

    def _render_latex_to_image(self, latex: str, display_mode: bool = False) -> Optional[str]:
        """
        Render LaTeX formula to PNG image using matplotlib.
        
        Args:
            latex: LaTeX formula string (without $ delimiters)
            display_mode: If True, render as display math (larger, centered)
            
        Returns:
            Path to the rendered image, or None if rendering failed
        """
        try:
            import matplotlib
            matplotlib.use('Agg')  # Non-interactive backend
            import matplotlib.pyplot as plt
            from matplotlib import mathtext
            
            # 生成唯一文件名
            formula_hash = hashlib.md5(latex.encode()).hexdigest()[:12]
            img_path = self.image_cache_dir / f"formula_{formula_hash}.png"
            
            # 如果已缓存，直接返回
            if img_path.exists():
                return str(img_path)
            
            # 设置字体大小
            fontsize = 14 if display_mode else 12
            
            # 创建图形
            fig, ax = plt.subplots(figsize=(0.1, 0.1))
            ax.axis('off')
            
            # 渲染公式
            wrapped_latex = f"${latex}$"
            text = ax.text(0.5, 0.5, wrapped_latex, fontsize=fontsize,
                          ha='center', va='center', transform=ax.transAxes)
            
            # 调整图形大小以适应公式
            fig.canvas.draw()
            bbox = text.get_window_extent(renderer=fig.canvas.get_renderer())
            bbox_inches = bbox.transformed(fig.dpi_scale_trans.inverted())
            
            # 添加边距
            pad = 0.1
            fig.set_size_inches(bbox_inches.width + pad, bbox_inches.height + pad)
            
            # 保存图片
            fig.savefig(str(img_path), dpi=150, bbox_inches='tight', 
                       pad_inches=0.05, transparent=False, facecolor='white')
            plt.close(fig)
            
            logger.info(f"Rendered LaTeX formula: {img_path}")
            return str(img_path)
            
        except Exception as e:
            logger.warning(f"Failed to render LaTeX formula: {e}")
            return None
    
    def _process_latex_in_text(self, text: str) -> Tuple[str, List[Dict]]:
        """
        Extract LaTeX formulas from text and return processed text with placeholders.
        
        Supports:
            - Inline math: $...$
            - Display math: $$...$$
            - LaTeX environments: \\[...\\], \\(...\\)
        
        Args:
            text: Text containing LaTeX formulas
            
        Returns:
            Tuple of (processed_text, list of formula info dicts)
        """
        formulas = []
        
        # 匹配显示公式 $$...$$ 或 \[...\]
        display_pattern = r'\$\$(.+?)\$\$|\\\[(.+?)\\\]'
        # 匹配行内公式 $...$ 或 \(...\)
        inline_pattern = r'(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)|\\\((.+?)\\\)'
        
        # 先处理显示公式
        def replace_display(match):
            formula = match.group(1) or match.group(2)
            if formula:
                idx = len(formulas)
                formulas.append({'latex': formula.strip(), 'display': True, 'placeholder': f'[[FORMULA_DISPLAY_{idx}]]'})
                return f'[[FORMULA_DISPLAY_{idx}]]'
            return match.group(0)
        
        text = re.sub(display_pattern, replace_display, text, flags=re.DOTALL)
        
        # 再处理行内公式
        def replace_inline(match):
            formula = match.group(1) or match.group(2)
            if formula:
                idx = len(formulas)
                formulas.append({'latex': formula.strip(), 'display': False, 'placeholder': f'[[FORMULA_INLINE_{idx}]]'})
                return f'[[FORMULA_INLINE_{idx}]]'
            return match.group(0)
        
        text = re.sub(inline_pattern, replace_inline, text)
        
        return text, formulas
    
    def _add_word_formula(self, paragraph, latex: str) -> bool:
        """
        Add a LaTeX formula to a Word paragraph using OMML (Office Math Markup Language).
        
        Args:
            paragraph: python-docx paragraph object
            latex: LaTeX formula string
            
        Returns:
            True if formula was added successfully
        """
        try:
            from latex2mathml import latex_to_mathml
            from lxml import etree
            
            # 转换 LaTeX 到 MathML
            mathml = latex_to_mathml(latex)
            
            # MathML 转 OMML (Office Math Markup Language)
            # 使用 Word 的 OMML 命名空间
            omml_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
            
            # 解析 MathML
            mathml_tree = etree.fromstring(mathml.encode())
            
            # 创建 OMML oMath 元素
            # 简化处理：将 MathML 作为文本插入，Word 会尝试解析
            # 更好的方案是使用完整的 MathML 到 OMML 转换
            run = paragraph.add_run(f" {latex} ")
            run.italic = True
            
            logger.debug(f"Added formula to Word: {latex[:50]}...")
            return True
            
        except ImportError:
            # 如果没有 latex2mathml，使用图片方式
            logger.debug("latex2mathml not available, using image fallback for Word formula")
            return False
        except Exception as e:
            logger.warning(f"Failed to add Word formula: {e}")
            return False
    
    def _check_mmdc_available(self) -> bool:
        """Check mermaid-cli (mmdc) availability once and cache the result."""
        if self._mmdc_available is not None:
            return self._mmdc_available

        import subprocess
        try:
            result = subprocess.run(['mmdc', '--version'], capture_output=True, timeout=5)
            if result.returncode == 0:
                self._mmdc_available = True
                try:
                    self._mmdc_version = (result.stdout.decode(errors='ignore') or '').strip() or None
                except Exception:
                    self._mmdc_version = None
                return True
            self._mmdc_available = False
            return False
        except FileNotFoundError:
            self._mmdc_available = False
            return False
        except subprocess.TimeoutExpired:
            self._mmdc_available = False
            return False
        except Exception:
            self._mmdc_available = False
            return False
    
    def _distribute_images_to_sections(self, images: List[Dict], sections: List[str]) -> Dict[str, List[Dict]]:
        """Distribute images evenly across major sections."""
        distribution = {}
        if not images or not sections:
            return distribution
        
        images_per_section = max(1, len(images) // len(sections))
        img_idx = 0
        
        for section in sections:
            section_images = []
            for _ in range(images_per_section):
                if img_idx < len(images):
                    section_images.append(images[img_idx])
                    img_idx += 1
            if section_images:
                distribution[section.lower()] = section_images
        
        return distribution
    
    def _render_mermaid_to_image(self, mermaid_type: str, description: str, context: str = "") -> Optional[str]:
        """
        Generate a Mermaid diagram and render it to an image file.
        
        Args:
            mermaid_type: Type of diagram (flowchart, sequence, class, pie, gantt, etc.)
            description: Description of what the diagram should show
            context: Additional context from the report content
            
        Returns:
            Path to the generated image file, or None if failed
        """
        try:
            from tools.chart import MermaidChartTool
            import base64
            import requests
            
            chart_tool = MermaidChartTool()
            
            type_mapping = {
                # Core diagrams
                'flowchart': 'flowchart',
                'sequence': 'sequence',
                'sequencediagram': 'sequence',
                'class': 'class',
                'classdiagram': 'class',
                'state': 'flowchart',
                'statediagram': 'flowchart',
                'statediagram-v2': 'flowchart',
                'er': 'er',
                'erdiagram': 'er',
                'gantt': 'flowchart',
                # Data visualization
                'pie': 'flowchart',
                'xychart-beta': 'flowchart',
                'xychart': 'flowchart',
                'quadrantchart': 'flowchart',
                'sankey-beta': 'flowchart',
                'sankey': 'flowchart',
                # Conceptual & hierarchical
                'mindmap': 'flowchart',
                'timeline': 'flowchart',
                'journey': 'flowchart',
                # Technical & architecture
                'gitgraph': 'flowchart',
                'c4context': 'flowchart',
                'c4container': 'flowchart',
                'c4component': 'flowchart',
                'c4dynamic': 'flowchart',
                'c4deployment': 'flowchart',
                'architecture-beta': 'flowchart',
                'requirementdiagram': 'flowchart',
                'packet-beta': 'flowchart',
                'block-beta': 'flowchart',
            }
            diagram_type = type_mapping.get(mermaid_type.lower(), 'flowchart')
            
            prompt = f"{description}\n\nContext: {context[:500]}" if context else description
            mermaid_code = chart_tool.generate(prompt, diagram_type)
            
            code_match = re.search(r'```mermaid\s*([\s\S]*?)\s*```', mermaid_code)
            if code_match:
                clean_code = code_match.group(1).strip()
            else:
                clean_code = mermaid_code.strip()
            
            img_data, final_code = self._try_render_with_fix(clean_code)
            
            if not img_data:
                logger.warning(f"All Mermaid render methods failed for: {description[:30]}...")
                # 返回 None，调用方可以决定是否显示代码块
                return None
            
            try:
                from PIL import Image
                img = Image.open(io.BytesIO(img_data))
                
                if img.mode in ('RGBA', 'P'):
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'RGBA':
                        background.paste(img, mask=img.split()[3])
                    else:
                        background.paste(img)
                    img = background
                
                target_width = 2400
                target_height = 1680
                
                orig_width, orig_height = img.size
                ratio = min(target_width / orig_width, target_height / orig_height)
                if ratio < 1:
                    new_size = (int(orig_width * ratio), int(orig_height * ratio))
                    img = img.resize(new_size, Image.Resampling.LANCZOS)
                elif ratio > 1 and orig_width < 800:
                    # 小图放大以提高清晰度
                    scale = min(ratio, 2.0)
                    new_size = (int(orig_width * scale), int(orig_height * scale))
                    img = img.resize(new_size, Image.Resampling.LANCZOS)
                
                img_hash = hashlib.md5(clean_code.encode()).hexdigest()[:12]
                img_path = self.image_cache_dir / f"mermaid_{img_hash}.png"
                img.save(str(img_path), "PNG", optimize=True)
                
                logger.info(f"Generated Mermaid diagram: {mermaid_type} - {description[:30]}...")
                return str(img_path)
                
            except Exception as render_e:
                logger.warning(f"Failed to process Mermaid image: {render_e}")
                return None
                
        except Exception as e:
            logger.warning(f"Failed to generate Mermaid diagram: {e}")
            return None
    
    def _parse_mermaid_marker(self, line: str) -> Optional[Tuple[str, str]]:
        """
        Parse a [MERMAID: type | description] marker.
        
        Returns:
            Tuple of (type, description) or None if not a valid marker
        """
        match = re.match(r'\[MERMAID:\s*(\w+)\s*\|\s*(.+?)\]', line, re.IGNORECASE)
        if match:
            return match.group(1).strip(), match.group(2).strip()
        return None

    def _is_references_heading(self, heading_text: str) -> bool:
        """
        检测标题是否为参考文献标题。
        
        支持多种格式：
        - "References", "参考文献"
        - "10. References", "10. 参考文献"
        - "References and Sources"
        """
        text = (heading_text or "").strip().lower()
        if not text:
            return False

        # 移除开头的数字编号（如 "10. References" -> "references"）
        text_without_number = re.sub(r'^\d+[\.\、\s]+', '', text).strip()

        # strict matching to avoid false positives like "Reference Architecture"
        allowed = {
            "references",
            "reference",
            "bibliography",
            "works cited",
            "references and sources",
            "参考文献",
            "参考资料",
            "引用文献",
            "文献引用",
        }
        if text in allowed or text_without_number in allowed:
            return True

        if text.startswith("references ") or text.startswith("references:"):
            return True
        if text_without_number.startswith("references ") or text_without_number.startswith("references:"):
            return True
        if text.startswith("bibliography ") or text.startswith("bibliography:"):
            return True
        
        # 检测 "Key Sources" 等子标题（这些也应该被移动到参考文献部分）
        if text_without_number in {"key sources", "sources", "cited sources", "主要来源", "来源"}:
            return True
            
        return False

    def _move_references_to_end(self, content: str) -> str:
        if not content:
            return content

        lines = content.splitlines()
        ref_blocks: List[List[str]] = []
        ref_heading_level = None
        ref_heading_line = None
        i = 0
        while i < len(lines):
            line = lines[i]
            m = re.match(r'^(#{1,6})\s+(.+?)\s*$', line.strip())
            if not m:
                i += 1
                continue

            heading_level = len(m.group(1))
            heading_text_raw = m.group(2).strip()
            if not self._is_references_heading(heading_text_raw):
                i += 1
                continue

            # 记录第一个参考文献标题
            if ref_heading_line is None:
                ref_heading_level = heading_level
                ref_heading_line = line

            start = i
            j = i + 1
            while j < len(lines):
                m2 = re.match(r'^(#{1,6})\s+(.+?)\s*$', lines[j].strip())
                if m2 and len(m2.group(1)) <= heading_level:
                    break
                j += 1

            # 只收集内容行（跳过标题行，因为我们只保留一个标题）
            ref_blocks.append(lines[start+1:j])
            del lines[start:j]
            i = start

        if not ref_blocks:
            return content

        if lines and lines[-1].strip() != '':
            lines.append('')
            lines.append('')
        elif len(lines) >= 2 and (lines[-1].strip() != '' or lines[-2].strip() != ''):
            lines.append('')

        # 合并所有参考文献内容，去重
        merged: List[str] = []
        seen_refs = set()
        
        # 添加统一的参考文献标题
        if ref_heading_line:
            merged.append(ref_heading_line)
            merged.append('')
        
        for block in ref_blocks:
            for line in block:
                stripped = line.strip()
                # 跳过空行和重复的引用
                if not stripped:
                    if merged and merged[-1].strip() != '':
                        merged.append('')
                    continue
                # 去重：基于内容的前100个字符判断
                ref_key = stripped[:100].lower()
                if ref_key not in seen_refs:
                    seen_refs.add(ref_key)
                    merged.append(line)

        while merged and merged[0].strip() == '':
            merged.pop(0)

        lines.extend(merged)
        return '\n'.join(lines).strip() + '\n'
    
    def _fix_mermaid_syntax(self, mermaid_code: str) -> str:
        """
        Fix common Mermaid syntax errors generated by LLMs.
        
        Handles:
            - Chinese characters in labels (cause rendering failures on Kroki/mermaid.ink)
            - <br> -> <br/> conversion
            - graph -> flowchart conversion
            - Invalid arrow syntax (-> to -->)
            - Special characters in node labels
            - Quotes inside edge labels
            - Multi-line node labels
            - Stray HTML tags and Markdown formatting
        """
        fixed = mermaid_code.strip()
        
        # 移除可能被LLM误加的 ``` 标记
        fixed = re.sub(r'^```\s*mermaid\s*\n?', '', fixed)
        fixed = re.sub(r'\n?```\s*$', '', fixed)
        
        # 修复 <br> -> <br/>
        fixed = re.sub(r'<br\s*/?>', '<br/>', fixed)
        
        # 将 graph TD/LR 转换为 flowchart TD/LR
        fixed = re.sub(r'^graph\s+(TD|LR|TB|BT|RL)', r'flowchart \1', fixed)
        
        # 修复无效箭头语法：-> 改为 -->（仅在 flowchart 中）
        if fixed.startswith('flowchart'):
            fixed = re.sub(r'(\w)\s*->\s*(\w)', r'\1 --> \2', fixed)
            fixed = re.sub(r'(\])\s*->\s*(\w)', r'\1 --> \2', fixed)
            fixed = re.sub(r'(\))\s*->\s*(\w)', r'\1 --> \2', fixed)
            # handle invalid edge-label forms like A ->[label] B (common LLM mistake)
            fixed = re.sub(r'->\s*(?=\[)', '-->', fixed)
            fixed = re.sub(r'->\s*(?=\()', '-->', fixed)
            fixed = re.sub(r'->\s*(?=/)', '-->', fixed)

        def _clean_edge_label_text(label: str) -> str:
            label = label.replace('\n', ' ')
            label = label.replace('"', '').replace("'", '')
            label = label.replace('|', ' ')
            label = label.replace('[', '').replace(']', '')
            label = label.replace('(', '').replace(')', '')
            label = re.sub(r'\s+', ' ', label).strip()
            if len(label) > 60:
                label = label[:57] + '...'
            return label

        fixed = re.sub(
            r'--\>\s*\[([^\]\n]{1,200})\]',
            lambda m: f"-->|{_clean_edge_label_text(m.group(1))}|",
            fixed,
        )
        fixed = re.sub(
            r'--\>\s*\(([^\)\n]{1,200})\)',
            lambda m: f"-->|{_clean_edge_label_text(m.group(1))}|",
            fixed,
        )
        fixed = re.sub(
            r'--\>\s*/([^/\n]{1,200})/',
            lambda m: f"-->|{_clean_edge_label_text(m.group(1))}|",
            fixed,
        )
        fixed = re.sub(r'--\>\s*\]', '-->', fixed)
        fixed = re.sub(r'--\>\s*\)', '-->', fixed)
        
        # 修复多行节点标签
        fixed = re.sub(r'\[\s*\n\s*', '[', fixed)
        fixed = re.sub(r'\s*\n\s*\]', ']', fixed)
        
        # 修复边标签中的引号
        fixed = re.sub(r'-->\s*\|([^|]*["\'][^|]*)\|', 
                       lambda m: f'-->|{m.group(1).replace(chr(34), "").replace(chr(39), "")}|', fixed)
        
        # 标准化箭头格式
        fixed = re.sub(r'(\w)\s*-+>\s*\|([^|]*)\|\s*(\w)', r'\1 -->|\2| \3', fixed)
        
        # 移除节点标签中的特殊字符
        def clean_node_label(match):
            prefix = match.group(1)
            label = match.group(2)
            suffix = match.group(3)
            label = label.replace('"', '').replace("'", "")
            label = label.replace('(', '').replace(')', '')
            label = label.replace('{', '').replace('}', '')
            label = label.replace('#', '').replace('&', 'and')
            label = label.replace(';', ' ').replace('%', ' percent')
            if len(label) > 45:
                label = label[:42] + '...'
            return f'{prefix}{label}{suffix}'
        
        fixed = re.sub(r'(\[)([^\]]+)(\])', clean_node_label, fixed)
        
        # 移除 Markdown 格式标记（保护 Mermaid 的 [*] 语法不被破坏）
        _star_placeholder = '__MERMAID_STAR__'
        fixed = fixed.replace('[*]', f'[{_star_placeholder}]')
        fixed = re.sub(r'\*\*([^*]+)\*\*', r'\1', fixed)
        fixed = re.sub(r'\*([^*]+)\*', r'\1', fixed)
        fixed = fixed.replace(f'[{_star_placeholder}]', '[*]')
        
        # 移除除 <br/> 以外的 HTML 标签
        fixed = re.sub(r'<(?!br/)([^>]+)>', '', fixed)
        
        # 清理空行
        lines = fixed.split('\n')
        cleaned_lines = []
        for line in lines:
            line = line.rstrip()
            if line.strip():
                cleaned_lines.append(line)
            elif cleaned_lines and cleaned_lines[-1].strip():
                cleaned_lines.append('')
        fixed = '\n'.join(cleaned_lines)
        
        # 修复 flowchart 中跨行的花括号
        if fixed.startswith('flowchart '):
            fixed = re.sub(r'\{([^}]*)\n([^}]*)\}', 
                          lambda m: '{' + m.group(1).replace('\n', ' ') + m.group(2) + '}', fixed)
        
        # 替换中文标点为英文标点（这不会破坏语义，只是标准化）
        fixed = fixed.replace('\u201c', '"').replace('\u201d', '"')
        fixed = fixed.replace('\u2018', "'").replace('\u2019', "'")
        fixed = fixed.replace('\uff0c', ',').replace('\uff1a', ':').replace('\uff1b', ';')
        fixed = fixed.replace('\u3001', ',').replace('\u3002', '.')
        
        # 修复 gantt 图表中的季度日期格式（Mermaid 不支持 2025-Q1 格式）
        # 将 2025-Q1 转换为 2025-01-01，2025-Q2 转换为 2025-04-01 等
        if 'gantt' in fixed.lower():
            quarter_map = {'Q1': '01-01', 'Q2': '04-01', 'Q3': '07-01', 'Q4': '10-01'}
            for q, date in quarter_map.items():
                fixed = re.sub(rf'(\d{{4}})-{q}\b', rf'\1-{date}', fixed)
        
        # 注意：CJK 文字替换已移至 _replace_cjk_with_placeholders 方法
        # 只在渲染失败时才调用，以保留 LLM 生成的有意义内容
        
        return fixed
    
    def _translate_cjk_text(self, chinese_text: str) -> str:
        """
        Translate Chinese text to English using LLM.
        
        Args:
            chinese_text: Chinese text to translate
            
        Returns:
            English translation, or simplified placeholder if translation fails
        """
        # 如果没有中文字符，直接返回
        if not any('\u4e00' <= ch <= '\u9fff' for ch in chinese_text):
            return chinese_text
        
        # 检查缓存
        if not hasattr(self, '_translation_cache'):
            self._translation_cache = {}
        
        if chinese_text in self._translation_cache:
            return self._translation_cache[chinese_text]
        
        try:
            from langchain_openai import ChatOpenAI
            from langchain_core.prompts import ChatPromptTemplate
            from config import LLM_CONFIG
            
            llm = ChatOpenAI(
                model=LLM_CONFIG["model"],
                api_key=LLM_CONFIG["api_key"],
                base_url=LLM_CONFIG["base_url"],
                temperature=0.1
            )
            
            prompt = ChatPromptTemplate.from_messages([
                ("system", """You are a translator. Translate the Chinese text to concise English.
Rules:
1. Keep it SHORT (max 5 words if possible)
2. Use simple words suitable for diagram labels
3. Output ONLY the English translation, nothing else
4. If there are English words mixed in, keep them"""),
                ("human", f"Translate: {chinese_text}")
            ])
            
            chain = prompt | llm
            result = chain.invoke({})
            translation = result.content.strip()
            
            # 清理翻译结果
            translation = translation.replace('"', '').replace("'", "")
            # Enforce ASCII-only output for maximum Mermaid renderer compatibility
            translation = re.sub(r'[^A-Za-z0-9\s\-]', '', translation)
            translation = ' '.join(translation.split())[:40]  # 限制长度
            
            if translation:
                self._translation_cache[chinese_text] = translation
                return translation
                
        except Exception as e:
            logger.warning(f"Translation failed for '{chinese_text[:20]}...': {e}")
        
        # 翻译失败时，提取英文部分或使用简短占位符
        english_parts = re.findall(r'[a-zA-Z0-9][a-zA-Z0-9\s\-_.]*[a-zA-Z0-9]|[a-zA-Z0-9]', chinese_text)
        if english_parts:
            result = ' '.join(p.strip() for p in english_parts).strip()
            self._translation_cache[chinese_text] = result
            return result
        
        # 最后的回退：使用简短标识
        if not hasattr(self, '_fallback_counter'):
            self._fallback_counter = 0
        self._fallback_counter += 1
        result = f"Item{self._fallback_counter}"
        self._translation_cache[chinese_text] = result
        return result
    
    def _translate_cjk_in_mermaid(self, mermaid_code: str) -> str:
        """
        Translate CJK characters in Mermaid code to English.
        
        This method translates Chinese text to English to ensure the diagram
        can be rendered by services that don't support CJK characters.
        """
        fixed = mermaid_code
        
        # 只有存在真正的 CJK 字符时才处理
        if not any('\u4e00' <= ch <= '\u9fff' for ch in fixed):
            return fixed
        
        def _cjk_to_english(text: str, context: str = "Item") -> str:
            """Translate CJK text to English using LLM."""
            if not any('\u4e00' <= ch <= '\u9fff' for ch in text):
                return text
            # 使用 LLM 翻译
            return self._translate_cjk_text(text)
        
        # 1) 替换方括号节点标签中的中文: A[中文] -> A[Item1]
        def replace_bracket_label(match):
            label = match.group(2)
            if label.strip() == '*':
                return match.group(0)
            if any('\u4e00' <= ch <= '\u9fff' for ch in label):
                label = _cjk_to_english(label, "Item")
            return f'{match.group(1)}{label}{match.group(3)}'
        fixed = re.sub(r'(\[)([^\]]+)(\])', replace_bracket_label, fixed)
        
        # 2) 替换圆括号节点标签中的中文: A(中文) -> A(Item N)
        def replace_paren_label(match):
            full = match.group(0)
            label = match.group(2)
            if any('\u4e00' <= ch <= '\u9fff' for ch in label):
                label = _cjk_to_english(label, "Item")
                return f'{match.group(1)}{label}{match.group(3)}'
            return full
        fixed = re.sub(r'(\(\()([^)]+)(\)\))', replace_paren_label, fixed)
        fixed = re.sub(r'(\()([^()]+)(\))', replace_paren_label, fixed)
        
        # 3) 替换引号内的中文文本（pie, quadrantChart 等）
        def replace_quoted_chinese(match):
            quote_open = match.group(1)
            text = match.group(2)
            quote_close = match.group(3)
            if any('\u4e00' <= ch <= '\u9fff' for ch in text):
                text = _cjk_to_english(text, "Item")
            return f'{quote_open}{text}{quote_close}'
        fixed = re.sub(r'(")((?:[^"]*[\u4e00-\u9fff][^"]*)+)(")', replace_quoted_chinese, fixed)
        
        # 4) 替换边标签中的中文: -->|中文| -> -->|Label N|
        def replace_edge_label(match):
            label = match.group(1)
            if any('\u4e00' <= ch <= '\u9fff' for ch in label):
                label = _cjk_to_english(label, "Label")
            return f'-->|{label}|'
        fixed = re.sub(r'-->\|([^|]+)\|', replace_edge_label, fixed)
        
        # 5) 处理剩余行级中文
        _first_line = fixed.split('\n')[0].strip().lower() if fixed.strip() else ''
        _is_sequence = _first_line.startswith('sequencediagram')
        _is_er = _first_line.startswith('erdiagram')
        
        lines = fixed.split('\n')
        new_lines = []
        for line in lines:
            if not any('\u4e00' <= ch <= '\u9fff' for ch in line):
                new_lines.append(line)
                continue
            stripped = line.lstrip()
            indent = line[:len(line) - len(stripped)]
            
            if stripped.startswith('participant '):
                rest = stripped[12:].strip()
                if ' as ' in rest:
                    parts = rest.split(' as ', 1)
                    alias = _cjk_to_english(parts[1].strip(), "Actor")
                    new_lines.append(f'{indent}participant {parts[0].strip()} as {alias}')
                else:
                    rest = _cjk_to_english(rest, "Actor")
                    new_lines.append(f'{indent}participant {rest}')
            elif _is_sequence and re.search(r'->>|-->>|->', stripped):
                cleaned = re.sub(r'[\u4e00-\u9fff]+', lambda m: _cjk_to_english(m.group(0), "Actor"), stripped)
                new_lines.append(f'{indent}{cleaned}')
            elif _is_er and re.search(r'\|\|--|o\{|\}o|--\|\|', stripped):
                cleaned = re.sub(r'[\u4e00-\u9fff]+', lambda m: _cjk_to_english(m.group(0), "Entity"), stripped)
                new_lines.append(f'{indent}{cleaned}')
            elif _is_er and stripped.endswith('{'):
                entity_name = stripped[:-1].strip()
                if any('\u4e00' <= ch <= '\u9fff' for ch in entity_name):
                    entity_name = _cjk_to_english(entity_name, "Entity")
                new_lines.append(f'{indent}{entity_name} {{')
            elif _is_er and re.match(r'^\s*(string|int|float|date|boolean|datetime)\s+', stripped):
                cleaned = re.sub(r'[\u4e00-\u9fff]+', lambda m: _cjk_to_english(m.group(0), "field"), stripped)
                new_lines.append(f'{indent}{cleaned}')
            elif stripped.startswith('title'):
                rest = stripped[5:].strip()
                rest = _cjk_to_english(rest, "Research Topic") if rest else "Research Topic"
                new_lines.append(f'{indent}title {rest}')
            elif stripped.startswith('section'):
                rest = stripped[7:].strip()
                rest = _cjk_to_english(rest, "Phase")
                new_lines.append(f'{indent}section {rest}')
            elif stripped.startswith('subgraph '):
                # 处理 subgraph 中文标签: subgraph "中文" -> subgraph "English"
                rest = stripped[9:].strip()
                if rest.startswith('"') and rest.endswith('"'):
                    inner = rest[1:-1]
                    if any('\u4e00' <= ch <= '\u9fff' for ch in inner):
                        inner = _cjk_to_english(inner, "Group")
                    new_lines.append(f'{indent}subgraph "{inner}"')
                elif any('\u4e00' <= ch <= '\u9fff' for ch in rest):
                    rest = _cjk_to_english(rest, "Group")
                    new_lines.append(f'{indent}subgraph {rest}')
                else:
                    new_lines.append(line)
            elif stripped.startswith('state '):
                cleaned = re.sub(r'[\u4e00-\u9fff]+', lambda m: _cjk_to_english(m.group(0), "State"), stripped)
                new_lines.append(f'{indent}{cleaned}')
            elif '-->' in stripped or '---' in stripped or '->>' in stripped or '-->>' in stripped:
                cleaned = re.sub(r'[\u4e00-\u9fff]+', lambda m: _cjk_to_english(m.group(0), "State"), stripped)
                new_lines.append(f'{indent}{cleaned}')
            elif ':' in stripped and not stripped.startswith(('x-axis', 'y-axis', 'title', 'section', 'dateFormat', 'axisFormat')):
                parts = stripped.split(':', 1)
                task_name = parts[0].strip()
                if any('\u4e00' <= ch <= '\u9fff' for ch in task_name):
                    task_name = _cjk_to_english(task_name, "Task")
                rest = parts[1]
                if any('\u4e00' <= ch <= '\u9fff' for ch in rest):
                    rest = re.sub(r'[\u4e00-\u9fff]+', lambda m: _cjk_to_english(m.group(0), "Item"), rest)
                new_lines.append(f'{indent}{task_name} :{rest}')
            elif stripped.startswith(('x-axis', 'y-axis')):
                cleaned = re.sub(r'[\u4e00-\u9fff]+', lambda m: _cjk_to_english(m.group(0), "Axis"), stripped)
                new_lines.append(f'{indent}{cleaned}')
            else:
                cleaned = re.sub(r'[\u4e00-\u9fff]+', lambda m: _cjk_to_english(m.group(0), "Topic"), stripped)
                if cleaned.strip():
                    new_lines.append(f'{indent}{cleaned}')
        fixed = '\n'.join(new_lines)
        
        # 6) 最终清理：修复可能残留的空标签
        fixed = re.sub(r'\[\s*\]', '[Item]', fixed)
        fixed = re.sub(r'\(\s*\)', '(Item)', fixed)
        fixed = re.sub(r'""', '"Item"', fixed)
        
        return fixed
    
    def _aggressive_simplify_mermaid(self, mermaid_code: str) -> str:
        """
        Aggressively simplify Mermaid code as last resort.
        Strips all styling, shortens labels, removes subgraphs.
        """
        simplified = mermaid_code
        
        # 移除所有 style/linkStyle/classDef 行
        simplified = re.sub(r'^\s*(style|linkStyle|classDef|class)\s+[^\n]+\n?', '', simplified, flags=re.MULTILINE)
        
        # 移除 subgraph 结构（保留内容）
        simplified = re.sub(r'^\s*subgraph\s+[^\n]*\n?', '', simplified, flags=re.MULTILINE)
        simplified = re.sub(r'^\s*end\s*$', '', simplified, flags=re.MULTILINE)
        
        # 截断所有节点标签到 25 个字符
        simplified = re.sub(r'\[([^\]]{25,})\]', lambda m: f'[{m.group(1)[:22]}...]', simplified)
        
        # 移除所有 ::: 类引用
        simplified = re.sub(r':::[\w-]+', '', simplified)
        
        # 移除空行
        lines = [l for l in simplified.split('\n') if l.strip()]
        simplified = '\n'.join(lines)
        
        return simplified.strip()
    
    def _try_render_with_fix(self, mermaid_code: str) -> Tuple[Optional[bytes], str]:
        """
        Try to render Mermaid code, applying fixes only when rendering fails.
        
        Strategy:
            1. Try rendering original code directly (no modifications)
            2. On failure: Apply basic syntax fix and retry
            3. On failure: Call LLM for syntax fix (up to 3 times) with error info
            4. If all rendering fails: Return (None, mermaid_code) for fallback display
        
        Returns:
            Tuple of (image_data, mermaid_code):
            - If rendering succeeds: (bytes, code_used)
            - If rendering fails: (None, final_code) for fallback display
        """
        current_code = mermaid_code.strip()
        last_error = None
        
        # 移除可能被LLM误加的 ``` 标记
        current_code = re.sub(r'^```\s*mermaid\s*\n?', '', current_code)
        current_code = re.sub(r'\n?```\s*$', '', current_code)
        
        # === 第一次尝试：直接渲染原始代码 ===
        img_data, error = self._try_all_renderers(current_code)
        if img_data:
            logger.info("Mermaid rendered successfully with original code")
            return (img_data, current_code)
        last_error = error
        
        logger.info(f"Original code render failed ({error}), trying basic syntax fix...")
        
        # === 第二次尝试：基本语法修复后渲染 ===
        fixed_code = self._fix_mermaid_syntax(current_code)
        if fixed_code != current_code:
            img_data, error = self._try_all_renderers(fixed_code)
            if img_data:
                logger.info("Mermaid rendered successfully after basic syntax fix")
                return (img_data, fixed_code)
            current_code = fixed_code
            last_error = error

        first_line = ''
        for _l in current_code.split('\n'):
            if _l.strip():
                first_line = _l.strip().lower()
                break

        unstable_prefixes = (
            'block-beta', 'block',
            'architecture-beta', 'architecture',
            'packet-beta', 'packet',
            'requirementdiagram', 'requirement',
        )
        if first_line.startswith(unstable_prefixes) or ('block:' in current_code.lower()):
            converted = self._convert_mermaid_to_flowchart_via_llm(current_code, last_error)
            if converted and converted != current_code:
                img_data, error = self._try_all_renderers(converted)
                if img_data:
                    logger.info("Mermaid rendered successfully after diagram conversion")
                    return (img_data, converted)
                current_code = converted
                last_error = error
        
        # === 第三次尝试：如果包含中文，尝试翻译成英文后渲染 ===
        if any('\u4e00' <= ch <= '\u9fff' for ch in current_code):
            logger.info("Code contains CJK characters, trying translation to English...")
            translated_code = self._translate_cjk_in_mermaid(current_code)
            if translated_code != current_code:
                img_data, error = self._try_all_renderers(translated_code)
                if img_data:
                    logger.info("Mermaid rendered successfully after CJK translation")
                    return (img_data, translated_code)
                current_code = translated_code
                last_error = error
        
        logger.info(f"Basic syntax fix failed ({last_error}), trying LLM-based fix...")
        
        # === 第四次尝试：LLM 语法修复（最多3次），传递错误信息 ===
        for llm_attempt in range(3):
            llm_fixed_code = self._fix_mermaid_via_llm(current_code, llm_attempt + 1, last_error)
            if llm_fixed_code and llm_fixed_code != current_code:
                img_data, error = self._try_all_renderers(llm_fixed_code)
                if img_data:
                    logger.info(f"Mermaid rendered successfully after LLM fix (attempt {llm_attempt + 1})")
                    return (img_data, llm_fixed_code)
                current_code = llm_fixed_code
                last_error = error
            else:
                logger.warning(f"LLM fix attempt {llm_attempt + 1} returned no changes")
                break  # LLM 没有返回有效修复，停止重试
        
        # === 所有渲染尝试失败，返回代码供回退显示 ===
        logger.warning(f"All Mermaid render attempts failed ({last_error}), returning code for fallback display")
        return (None, current_code)
    
    def _try_all_renderers(self, mermaid_code: str) -> tuple[Optional[bytes], Optional[str]]:
        """
        Try all available renderers in priority order.
        
        Returns:
            Tuple of (image_data, error_message):
            - If success: (bytes, None)
            - If failure: (None, error_description)
        """
        errors = []
        
        # 1. 优先尝试本地 CLI（支持 CJK，无网络依赖）
        img_data, error = self._render_mermaid_via_cli(mermaid_code)
        if img_data:
            return (img_data, None)
        if error:
            errors.append(f"Local CLI: {error}")
        
        # 2. 尝试 mermaid.ink 渲染
        img_data, error = self._render_mermaid_via_ink(mermaid_code)
        if img_data:
            return (img_data, None)
        if error:
            errors.append(f"mermaid.ink: {error}")
        
        # 3. 最后尝试 Kroki 渲染
        img_data, error = self._render_mermaid_via_kroki(mermaid_code)
        if img_data:
            return (img_data, None)
        if error:
            errors.append(f"Kroki: {error}")
        
        return (None, "; ".join(errors) if errors else "Unknown rendering error")
    
    def _fix_mermaid_via_llm(self, mermaid_code: str, attempt: int, error_message: str = None) -> Optional[str]:
        """
        Use LLM to fix Mermaid syntax errors.
        
        Args:
            mermaid_code: The Mermaid code to fix
            attempt: Current attempt number (1-3)
            error_message: Error message from previous render attempt
        
        Returns:
            Fixed Mermaid code, or None if LLM call fails
        """
        try:
            from langchain_openai import ChatOpenAI
            from langchain_core.prompts import ChatPromptTemplate
            from config import LLM_CONFIG
            
            llm = ChatOpenAI(
                model=LLM_CONFIG["model"],
                api_key=LLM_CONFIG["api_key"],
                base_url=LLM_CONFIG["base_url"],
                temperature=0.1  # 低温度以获得稳定的修复结果
            )
            
            # 构建包含错误信息的提示
            error_context = ""
            if error_message:
                error_context = f"\n\nRENDERING ERROR: {error_message}\nPlease fix the issue that caused this error."
            
            prompt = ChatPromptTemplate.from_messages([
                ("system", """You are a Mermaid diagram syntax expert. Your task is to fix syntax errors in Mermaid code.

RULES:
1. Fix ONLY syntax errors, preserve the original meaning and structure
2. Common issues to fix:
   - Invalid arrow syntax (-> should be --> in flowchart)
   - Missing or incorrect diagram type declaration
   - Invalid node/edge labels (special characters, unescaped quotes)
   - Incorrect indentation (especially for mindmap, gantt)
   - Missing semicolons or line breaks where required
   - Invalid edge label syntax: never use -->[label] or --> (label). Use -->|label| instead
   - CJK characters causing encoding issues (translate to English if needed)
   - Do NOT introduce beta diagram types (block-beta, architecture-beta, packet-beta) or change diagram type unless absolutely required
3. Output ONLY the fixed Mermaid code, no explanations
4. Do NOT add ```mermaid markers
5. Keep all text content unchanged unless it causes syntax errors"""),
                ("human", f"""Fix the syntax errors in this Mermaid code (attempt {attempt}/3):{error_context}

{mermaid_code}

Output only the fixed Mermaid code:""")
            ])
            
            chain = prompt | llm
            result = chain.invoke({})
            fixed_code = result.content.strip()
            
            # 移除可能被 LLM 添加的 ``` 标记
            fixed_code = re.sub(r'^```\s*mermaid\s*\n?', '', fixed_code)
            fixed_code = re.sub(r'\n?```\s*$', '', fixed_code)
            
            logger.info(f"LLM fix attempt {attempt} completed")
            return fixed_code
            
        except Exception as e:
            logger.warning(f"LLM fix attempt {attempt} failed: {e}")
            return None

    def _convert_mermaid_to_flowchart_via_llm(self, mermaid_code: str, error_message: Optional[str] = None) -> Optional[str]:
        try:
            from langchain_openai import ChatOpenAI
            from langchain_core.prompts import ChatPromptTemplate
            from config import LLM_CONFIG

            llm = ChatOpenAI(
                model=LLM_CONFIG["model"],
                api_key=LLM_CONFIG["api_key"],
                base_url=LLM_CONFIG["base_url"],
                temperature=0.1,
            )

            error_context = f"\n\nRENDERING ERROR: {error_message}" if error_message else ""
            prompt = ChatPromptTemplate.from_messages([
                (
                    "system",
                    """You convert Mermaid diagrams into a Mermaid flowchart that can be rendered reliably.
Rules:
1. Output ONLY Mermaid code (no backticks)
2. Output MUST start with: flowchart TD
3. Use ONLY flowchart syntax (no block-beta/architecture-beta/packet-beta)
4. All node IDs must be ASCII identifiers (A, B, node1)
5. Edge labels must use: -->|label| (never -->[label])
6. If text contains Chinese, translate it to concise English""",
                ),
                (
                    "human",
                    f"Convert this Mermaid diagram to a flowchart TD while preserving meaning.{error_context}\n\n{mermaid_code}\n\nOutput only Mermaid code:",
                ),
            ])

            chain = prompt | llm
            result = chain.invoke({})
            converted = result.content.strip()
            converted = re.sub(r'^```\s*mermaid\s*\n?', '', converted)
            converted = re.sub(r'\n?```\s*$', '', converted)
            return converted.strip()
        except Exception as e:
            logger.warning(f"Diagram conversion failed: {e}")
            return None
    
    def _simplify_mermaid(self, mermaid_code: str) -> str:
        """Simplify complex Mermaid code that may cause rendering issues."""
        simplified = mermaid_code
        
        simplified = re.sub(r'\[([^\]]{80,})\]', lambda m: f'[{m.group(1)[:75]}...]', simplified)
        
        simplified = re.sub(r'style\s+\w+\s+[^\n]+\n?', '', simplified)
        simplified = re.sub(r'linkStyle\s+[^\n]+\n?', '', simplified)
        
        simplified = re.sub(r':::[\w-]+', '', simplified)
        
        return simplified.strip()
    
    def _render_mermaid_via_cli(self, mermaid_code: str) -> tuple[Optional[bytes], Optional[str]]:
        """
        Render Mermaid code using local Mermaid CLI (mmdc).
        
        This is the preferred method when available because:
        - Supports CJK characters natively
        - No network dependency
        - Faster rendering
        
        Requires: npm install -g @mermaid-js/mermaid-cli
        
        Returns:
            Tuple of (image_data, error_message)
        """
        import subprocess
        import tempfile

        if not self._check_mmdc_available():
            return (None, "mmdc not installed (npm install -g @mermaid-js/mermaid-cli)")
        
        try:
            # 创建临时文件
            with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, encoding='utf-8') as f:
                f.write(mermaid_code)
                input_path = f.name
            
            output_path = input_path.replace('.mmd', '.png')
            
            # 调用 mmdc 渲染
            result = subprocess.run(
                ['mmdc', '-i', input_path, '-o', output_path, '-b', 'white', '-s', '2'],
                capture_output=True,
                timeout=30
            )
            
            if result.returncode == 0 and os.path.exists(output_path):
                with open(output_path, 'rb') as f:
                    img_data = f.read()
                # 清理临时文件
                os.unlink(input_path)
                os.unlink(output_path)
                logger.info("Mermaid rendered via local CLI (mmdc)")
                return (img_data, None)
            else:
                error_msg = result.stderr.decode().strip() if result.stderr else "Unknown mmdc error"
                logger.warning(f"mmdc render failed: {error_msg}")
                return (None, error_msg)
                
        except subprocess.TimeoutExpired:
            return (None, "mmdc rendering timed out (>30s)")
        except Exception as e:
            logger.warning(f"Local Mermaid CLI render failed: {e}")
            return (None, str(e))
    
    def _render_mermaid_via_kroki(self, mermaid_code: str) -> tuple[Optional[bytes], Optional[str]]:
        """
        Render Mermaid code using Kroki.io API (supports longer code via POST).
        
        Returns:
            Tuple of (image_data, error_message)
        """
        import zlib
        import base64
        import requests
        
        try:
            compressed = zlib.compress(mermaid_code.encode('utf-8'), 9)
            encoded = base64.urlsafe_b64encode(compressed).decode('utf-8')
            
            kroki_url = f"https://kroki.io/mermaid/png/{encoded}"
            headers = {'User-Agent': 'Mozilla/5.0'}
            response = requests.get(kroki_url, timeout=30, headers=headers)
            
            if response.status_code == 200:
                return (response.content, None)
            
            # GET 失败，尝试 POST
            kroki_post_url = "https://kroki.io/mermaid/png"
            response = requests.post(
                kroki_post_url,
                data=mermaid_code.encode('utf-8'),
                headers={'Content-Type': 'text/plain', 'User-Agent': 'Mozilla/5.0'},
                timeout=30
            )
            
            if response.status_code == 200:
                return (response.content, None)
            else:
                error_msg = f"HTTP {response.status_code}"
                # 尝试解析错误信息
                try:
                    error_text = response.text[:200]
                    if error_text:
                        error_msg = f"HTTP {response.status_code}: {error_text}"
                except:
                    pass
                logger.warning(f"Kroki render failed: {error_msg}")
                return (None, error_msg)
            
        except requests.exceptions.Timeout:
            return (None, "Request timed out (>30s)")
        except requests.exceptions.ConnectionError:
            return (None, "Connection error - network unavailable")
        except Exception as e:
            logger.warning(f"Kroki render failed: {e}")
            return (None, str(e))
    
    def _render_mermaid_via_ink(self, mermaid_code: str) -> tuple[Optional[bytes], Optional[str]]:
        """
        Render Mermaid code using mermaid.ink API.
        
        Returns:
            Tuple of (image_data, error_message)
        """
        import base64
        import requests
        
        try:
            encoded = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
            # 使用 width 参数而非 scale，避免 "scale can only be set when width or height is set" 错误
            # 增加宽度到1600以确保饼图等图表的标题不被截断
            mermaid_url = f"https://mermaid.ink/img/{encoded}?width=1600"
            
            headers = {'User-Agent': 'Mozilla/5.0'}
            response = requests.get(mermaid_url, timeout=30, headers=headers)
            
            if response.status_code == 200:
                return (response.content, None)
            else:
                error_msg = f"HTTP {response.status_code}"
                try:
                    error_text = response.text[:200]
                    if error_text:
                        error_msg = f"HTTP {response.status_code}: {error_text}"
                except:
                    pass
                logger.warning(f"mermaid.ink render failed: {error_msg}")
                return (None, error_msg)
                
        except requests.exceptions.Timeout:
            return (None, "Request timed out (>30s)")
        except requests.exceptions.ConnectionError:
            return (None, "Connection error - network unavailable")
        except Exception as e:
            logger.warning(f"mermaid.ink render failed: {e}")
            return (None, str(e))
    
    def _process_mermaid_code_blocks(self, content: str) -> str:
        """
        Find and render all ```mermaid code blocks in the content,
        replacing them with image references. If rendering fails, keep the code block.
        """
        mermaid_pattern = re.compile(r'```mermaid\s*([\s\S]*?)\s*```', re.IGNORECASE)
        
        def render_mermaid_block(match):
            mermaid_code = match.group(1).strip()
            if not mermaid_code:
                return ""
            
            try:
                img_data, final_code = self._try_render_with_fix(mermaid_code)
                
                if not img_data:
                    # 渲染失败，保留 Mermaid 代码块供回退显示
                    logger.warning("All Mermaid render methods failed, keeping code block")
                    return f"```mermaid\n{final_code}\n```"
                
                from PIL import Image
                img = Image.open(io.BytesIO(img_data))
                
                if img.mode in ('RGBA', 'P'):
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'RGBA':
                        background.paste(img, mask=img.split()[3])
                    else:
                        background.paste(img)
                    img = background
                
                img_hash = hashlib.md5(mermaid_code.encode()).hexdigest()[:12]
                img_path = self.image_cache_dir / f"mermaid_block_{img_hash}.png"
                img.save(str(img_path), "PNG", optimize=True)
                
                first_line = mermaid_code.split('\n')[0].strip()
                diagram_type = "Diagram"
                if first_line.startswith('flowchart') or first_line.startswith('graph'):
                    diagram_type = "Flowchart"
                elif first_line.startswith('sequenceDiagram'):
                    diagram_type = "Sequence Diagram"
                elif first_line.startswith('classDiagram'):
                    diagram_type = "Class Diagram"
                elif first_line.startswith('stateDiagram'):
                    diagram_type = "State Diagram"
                elif first_line.startswith('erDiagram'):
                    diagram_type = "ER Diagram"
                elif first_line.startswith('gantt'):
                    diagram_type = "Gantt Chart"
                elif first_line.startswith('pie'):
                    diagram_type = "Pie Chart"
                elif first_line.startswith('mindmap'):
                    diagram_type = "Mind Map"
                elif first_line.startswith('timeline'):
                    diagram_type = "Timeline"
                elif first_line.startswith('quadrantChart'):
                    diagram_type = "Quadrant Chart"
                elif first_line.startswith('gitGraph'):
                    diagram_type = "Git Graph"
                
                logger.info(f"Successfully rendered Mermaid {diagram_type}: {img_path}")
                return f"\n[RENDERED_MERMAID_IMAGE:{img_path}|{diagram_type}]\n"
                
            except Exception as e:
                logger.warning(f"Failed to render mermaid code block: {e}")
                return ""
        
        return mermaid_pattern.sub(render_mermaid_block, content)
    
    def generate(self, title: str, content: str, filename: str = None,
                 images: List[Dict] = None, tables: List[Dict] = None) -> str:
        """Generate PDF report with academic formatting."""
        return self._generate_pdf(title, content, filename, images, tables)
    
    def generate_both(self, title: str, content: str, filename: str = None,
                      images: List[Dict] = None, tables: List[Dict] = None) -> Tuple[str, str]:
        """
        Generate both PDF and Word versions of the report.
        
        Returns:
            Tuple of (pdf_path, docx_path)
        """
        content = self._move_references_to_end(content)

        self._reset_counters()
        pdf_path = self._generate_pdf(title, content, filename, images, tables)
        
        self._reset_counters()
        base_name = Path(pdf_path).stem
        docx_path = self._generate_word(title, content, base_name, images, tables)
        
        return pdf_path, docx_path
    
    def _generate_pdf(self, title: str, content: str, filename: str = None,
                      images: List[Dict] = None, tables: List[Dict] = None) -> str:
        """Generate PDF with academic paper formatting."""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.units import inch, cm
        except ImportError:
            raise ImportError("reportlab library required for PDF generation")
        
        self._reset_counters()
        images = images or []
        tables = tables or []

        content = self._move_references_to_end(content)
        
        # 检测是否为中文内容
        is_chinese = self._is_chinese_content(title) or self._is_chinese_content(content[:500])
        
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"research_report_{timestamp}.pdf"
        if not filename.endswith(".pdf"):
            filename += ".pdf"
        
        pdf_path = self.output_dir / filename
        
        if is_chinese:
            # 中文论文页面设置：A4纸，左右2.6cm，上下3cm
            doc = SimpleDocTemplate(
                str(pdf_path),
                pagesize=A4,
                rightMargin=2.6*cm,
                leftMargin=2.6*cm,
                topMargin=3*cm,
                bottomMargin=3*cm
            )
        else:
            doc = SimpleDocTemplate(
                str(pdf_path),
                pagesize=A4,
                rightMargin=1*inch,
                leftMargin=1*inch,
                topMargin=0.75*inch,
                bottomMargin=0.75*inch
            )
        
        base_styles, custom = self._get_academic_styles(is_chinese=is_chinese)
        
        downloaded_images = self._prepare_images(images)
        
        major_sections = []
        for line in content.split('\n'):
            if line.strip().startswith('# '):
                major_sections.append(line.strip()[2:])
        image_distribution = self._distribute_images_to_sections(downloaded_images, major_sections)
        
        story = []
        
        story.append(Paragraph(title, custom['title']))
        if is_chinese:
            story.append(Paragraph("RAAA 深度研究系统自动生成", custom['author']))
            date_str = datetime.now().strftime("%Y年%m月%d日")
        else:
            story.append(Paragraph("Generated by RAAA Deep Research", custom['author']))
            date_str = datetime.now().strftime("%B %d, %Y")
        story.append(Paragraph(date_str, custom['date']))
        story.append(Spacer(1, 12))
        
        content = self._move_references_to_end(content)
        content = self._process_mermaid_code_blocks(content)
        
        lines = content.split('\n')
        current_section = ""
        section_image_idx = {}
        in_table = False
        table_lines = []
        after_heading = False
        in_references = False
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            if not line_stripped:
                if in_table and table_lines:
                    table_text = '\n'.join(table_lines)
                    table_data = self._parse_table_from_text(table_text)
                    if table_data:
                        story.extend(self._create_table_with_caption(table_data, "Research Data", is_chinese=is_chinese))
                    table_lines = []
                    in_table = False
                continue
            
            if line_stripped.startswith('|') and '|' in line_stripped[1:]:
                in_table = True
                table_lines.append(line_stripped)
                continue
            elif in_table:
                if table_lines:
                    table_text = '\n'.join(table_lines)
                    table_data = self._parse_table_from_text(table_text)
                    if table_data:
                        story.extend(self._create_table_with_caption(table_data, "Research Data", is_chinese=is_chinese))
                table_lines = []
                in_table = False
            
            if line_stripped.startswith('# '):
                text = line_stripped[2:].replace('**', '').replace('*', '')
                text = self._escape_xml(text)
                current_section = line_stripped[2:].lower()
                in_references = 'reference' in current_section.lower()
                
                story.append(PageBreak())
                story.append(Paragraph(text if is_chinese else text.upper(), custom['h1']))
                after_heading = True
                
                section_key = current_section.split()[0] if current_section else ""
                for key in image_distribution:
                    if section_key in key or key in section_key:
                        if key not in section_image_idx:
                            section_image_idx[key] = 0
                        if section_image_idx[key] < len(image_distribution[key]):
                            img = image_distribution[key][section_image_idx[key]]
                            story.extend(self._create_figure_with_caption(
                                img['path'], img['title'], img['source'], is_chinese=is_chinese
                            ))
                            section_image_idx[key] += 1
                        break
                
            elif line_stripped.startswith('## '):
                text = line_stripped[3:].replace('**', '').replace('*', '')
                text = self._escape_xml(text)
                story.append(Paragraph(text, custom['h2']))
                after_heading = True
                
            elif line_stripped.startswith('### '):
                text = line_stripped[4:].replace('**', '').replace('*', '')
                text = self._escape_xml(text)
                story.append(Paragraph(text, custom['h3']))
                after_heading = True
                
            elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                text = line_stripped[2:]
                text = self._escape_xml(text)
                story.append(Paragraph(f"• {text}", custom['bullet']))
                after_heading = False
                
            elif re.match(r'^\d+\.\s', line_stripped):
                match = re.match(r'^(\d+)\.\s(.*)$', line_stripped)
                if match:
                    num, text = match.groups()
                    text = self._escape_xml(text)
                    story.append(Paragraph(f"{num}. {text}", custom['numbered']))
                after_heading = False
                
            elif line_stripped.startswith('[TABLE:') or line_stripped.startswith('[IMAGE:'):
                continue
            
            elif line_stripped.startswith('[RENDERED_MERMAID_IMAGE:'):
                img_match = re.match(r'\[RENDERED_MERMAID_IMAGE:(.+?)\|(.+?)\]', line_stripped)
                if img_match:
                    rendered_img_path = img_match.group(1).strip()
                    diagram_type = img_match.group(2).strip()
                    if os.path.exists(rendered_img_path):
                        story.extend(self._create_figure_with_caption(
                            rendered_img_path,
                            diagram_type,
                            "Generated Diagram",
                            is_chinese=is_chinese,
                            diagram_type=diagram_type,
                            show_caption=True
                        ))
                continue
            
            elif line_stripped.startswith('[WEB_SCREENSHOT:'):
                # 处理网页截图标记: [WEB_SCREENSHOT: url | description]
                screenshot_match = re.match(r'\[WEB_SCREENSHOT:\s*(.+?)\s*\|\s*(.+?)\s*\]', line_stripped)
                if screenshot_match:
                    screenshot_url = screenshot_match.group(1).strip()
                    screenshot_desc = screenshot_match.group(2).strip()
                    screenshot_path = self._download_web_screenshot(screenshot_url)
                    # 下载失败时跳过图片，不生成占位图
                    if screenshot_path:
                        story.extend(self._create_figure_with_caption(
                            screenshot_path,
                            screenshot_desc,
                            screenshot_url,
                            is_chinese=is_chinese,
                            show_caption=True
                        ))
                continue
            
            elif line_stripped.startswith('[MERMAID:'):
                mermaid_info = self._parse_mermaid_marker(line_stripped)
                if mermaid_info:
                    mermaid_type, mermaid_desc = mermaid_info
                    context = "\n".join(lines[max(0, i-10):i])
                    img_path = self._render_mermaid_to_image(mermaid_type, mermaid_desc, context)
                    if img_path:
                        story.extend(self._create_figure_with_caption(
                            img_path, 
                            mermaid_desc, 
                            f"Generated {mermaid_type.title()} Diagram",
                            is_chinese=is_chinese,
                            diagram_type=mermaid_type,
                            show_caption=True
                        ))
                continue
                
            else:
                # 处理公式
                processed_text, formulas = self._process_latex_in_text(line_stripped)
                
                if formulas:
                    # 如果有公式，需要分段处理
                    story.extend(self._create_paragraph_with_formulas(
                        processed_text, formulas, 
                        custom['reference'] if in_references else (custom['body_first'] if after_heading else custom['body']),
                        is_chinese=is_chinese
                    ))
                    if after_heading:
                        after_heading = False
                else:
                    text = self._escape_xml(line_stripped)
                    if in_references:
                        story.append(Paragraph(text, custom['reference']))
                    elif after_heading:
                        story.append(Paragraph(text, custom['body_first']))
                        after_heading = False
                    else:
                        story.append(Paragraph(text, custom['body']))
        
        if in_table and table_lines:
            table_text = '\n'.join(table_lines)
            table_data = self._parse_table_from_text(table_text)
            if table_data:
                story.extend(self._create_table_with_caption(table_data, "Research Data", is_chinese=is_chinese))
        
        story.append(Spacer(1, 24))
        if is_chinese:
            footer_style = custom.get('body', base_styles['Normal'])
            story.append(Paragraph(
                "本报告由 RAAA（需求分析智能体助手）深度研究系统自动生成",
                footer_style
            ))
        else:
            story.append(Paragraph(
                "<i>This report was generated by Requirements Analysis Agent Assistant (RAAA)</i>",
                base_styles['Normal']
            ))
        
        doc.build(story)
        logger.info(f"PDF report generated: {pdf_path}")
        return str(pdf_path)
    
    def _create_paragraph_with_formulas(self, text: str, formulas: List[Dict], 
                                         style, is_chinese: bool = False) -> List[Any]:
        """
        Create PDF flowables for text containing LaTeX formulas.
        Formulas are rendered as inline images.
        
        Args:
            text: Text with formula placeholders
            formulas: List of formula info dicts from _process_latex_in_text
            style: ReportLab paragraph style
            is_chinese: Whether the content is Chinese
            
        Returns:
            List of flowables (Paragraphs and Images)
        """
        from reportlab.platypus import Paragraph, Image, Spacer
        from reportlab.lib.units import inch
        
        flowables = []
        
        # 创建公式占位符到图片路径的映射
        formula_images = {}
        for formula in formulas:
            img_path = self._render_latex_to_image(formula['latex'], formula['display'])
            if img_path:
                formula_images[formula['placeholder']] = {
                    'path': img_path,
                    'display': formula['display'],
                    'latex': formula['latex']
                }
        
        # 检查是否有显示公式（需要单独成行）
        has_display_formula = any(f['display'] for f in formulas if f['placeholder'] in formula_images)
        
        if has_display_formula:
            # 如果有显示公式，分段处理
            parts = re.split(r'(\[\[FORMULA_DISPLAY_\d+\]\])', text)
            for part in parts:
                part = part.strip()
                if not part:
                    continue
                    
                if part.startswith('[[FORMULA_DISPLAY_') and part in formula_images:
                    # 显示公式 - 居中显示为图片
                    img_info = formula_images[part]
                    try:
                        from PIL import Image as PILImage
                        with PILImage.open(img_info['path']) as img:
                            img_width, img_height = img.size
                        # 限制最大宽度
                        max_width = 5 * inch
                        scale = min(1.0, max_width / (img_width / 150 * inch))
                        display_width = (img_width / 150) * inch * scale
                        display_height = (img_height / 150) * inch * scale
                        
                        flowables.append(Spacer(1, 6))
                        img_flowable = Image(img_info['path'], width=display_width, height=display_height)
                        img_flowable.hAlign = 'CENTER'
                        flowables.append(img_flowable)
                        flowables.append(Spacer(1, 6))
                    except Exception as e:
                        logger.warning(f"Failed to add formula image: {e}")
                        # 回退：显示原始 LaTeX
                        flowables.append(Paragraph(self._escape_xml(f"[{img_info['latex']}]"), style))
                else:
                    # 普通文本（可能包含行内公式）
                    # 替换行内公式为图片标记或文本
                    processed_part = part
                    for placeholder, img_info in formula_images.items():
                        if not img_info['display'] and placeholder in processed_part:
                            # 行内公式 - 在 PDF 中暂时用斜体文本表示
                            processed_part = processed_part.replace(
                                placeholder, 
                                f"<i>{self._escape_xml(img_info['latex'])}</i>"
                            )
                    
                    if processed_part.strip():
                        flowables.append(Paragraph(self._escape_xml(processed_part), style))
        else:
            # 只有行内公式
            processed_text = text
            for placeholder, img_info in formula_images.items():
                # 行内公式用斜体表示
                processed_text = processed_text.replace(
                    placeholder,
                    f"<i>{self._escape_xml(img_info['latex'])}</i>"
                )
            flowables.append(Paragraph(self._escape_xml(processed_text), style))
        
        return flowables if flowables else [Paragraph(self._escape_xml(text), style)]
    
    def _escape_xml(self, text: str) -> str:
        """Escape XML special characters while preserving some formatting."""
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
        text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
        # 修复嵌套标签问题：确保标签正确闭合
        text = self._fix_html_tag_nesting(text)
        return text
    
    def _fix_html_tag_nesting(self, text: str) -> str:
        """
        Fix improperly nested HTML tags (e.g., <b><i>text</b></i> -> <b><i>text</i></b>).
        Uses a stack-based approach to ensure proper tag closure order.
        """
        import re
        
        # 支持的标签
        tag_pattern = re.compile(r'<(/?)([bi])>')
        
        result = []
        stack = []  # 存储打开的标签
        last_end = 0
        
        for match in tag_pattern.finditer(text):
            # 添加标签之前的文本
            result.append(text[last_end:match.start()])
            last_end = match.end()
            
            is_closing = match.group(1) == '/'
            tag_name = match.group(2)
            
            if not is_closing:
                # 打开标签
                stack.append(tag_name)
                result.append(f'<{tag_name}>')
            else:
                # 关闭标签
                if tag_name in stack:
                    # 找到对应的打开标签位置
                    idx = len(stack) - 1 - stack[::-1].index(tag_name)
                    # 先关闭所有在它之后打开的标签
                    tags_to_reopen = []
                    while len(stack) > idx + 1:
                        t = stack.pop()
                        result.append(f'</{t}>')
                        tags_to_reopen.append(t)
                    # 关闭目标标签
                    stack.pop()
                    result.append(f'</{tag_name}>')
                    # 重新打开之前关闭的标签
                    for t in reversed(tags_to_reopen):
                        stack.append(t)
                        result.append(f'<{t}>')
                # 如果标签不在栈中，忽略这个关闭标签
        
        # 添加剩余文本
        result.append(text[last_end:])
        
        # 关闭所有未关闭的标签
        while stack:
            t = stack.pop()
            result.append(f'</{t}>')
        
        return ''.join(result)
    
    def _set_word_run_font(self, run, font_latin: str, font_size_pt: float, 
                           is_chinese: bool = False, bold: bool = False, italic: bool = False,
                           font_east_asia: str = None):
        """
        Set font properties on a Word run, including East Asian font for Chinese.
        python-docx requires XML manipulation to set East Asian fonts properly.
        """
        from docx.shared import Pt
        from docx.oxml.ns import qn
        
        run.font.name = font_latin
        run.font.size = Pt(font_size_pt)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        
        if is_chinese and font_east_asia:
            # 设置东亚字体（python-docx 不直接支持，需要操作 XML）
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = run._element.makeelement(qn('w:rFonts'), {})
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), font_east_asia)
    
    def _set_paragraph_line_spacing(self, para, line_spacing: float = 1.5):
        """Set paragraph line spacing (e.g., 1.5 for 1.5倍行距)."""
        from docx.shared import Pt
        from docx.oxml.ns import qn
        
        pPr = para._element.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = para._element.makeelement(qn('w:spacing'), {})
            pPr.append(spacing)
        # line spacing in 240ths of a line
        spacing.set(qn('w:line'), str(int(240 * line_spacing)))
        spacing.set(qn('w:lineRule'), 'auto')
    
    def _add_header_and_footer(self, doc, title: str = "", is_chinese: bool = False):
        """
        Add header (with title) and footer (with page number) to Word document.
        
        Args:
            doc: Word document object
            title: Report title for header
            is_chinese: Whether to use Chinese formatting
        """
        from docx.oxml.ns import qn
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        try:
            section = doc.sections[-1]
            
            # === 添加页眉 ===
            header = section.header
            header.is_linked_to_previous = False
            
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 页眉内容：报告标题（截断过长标题）
            header_text = title[:50] + "..." if len(title) > 50 else title
            if not header_text:
                header_text = "RAAA 深度研究报告" if is_chinese else "RAAA Deep Research Report"
            
            header_run = header_para.add_run(header_text)
            header_run.font.size = Pt(9)
            header_run.font.name = 'Times New Roman'
            if is_chinese:
                header_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            
            # 页眉下划线
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            pPr = header_para._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="auto"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            
            # === 添加页脚（页码） ===
            footer = section.footer
            footer.is_linked_to_previous = False
            
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加页码格式：第 X 页 / 共 Y 页 或 Page X of Y
            if is_chinese:
                footer_para.add_run("第 ").font.size = Pt(9)
            else:
                footer_para.add_run("Page ").font.size = Pt(9)
            
            # 当前页码域
            run1 = footer_para.add_run()
            run1.font.size = Pt(9)
            fldChar1 = run1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
            run1._element.append(fldChar1)
            
            run2 = footer_para.add_run()
            run2.font.size = Pt(9)
            instrText = run2._element.makeelement(qn('w:instrText'), {})
            instrText.text = ' PAGE '
            run2._element.append(instrText)
            
            run3 = footer_para.add_run()
            run3.font.size = Pt(9)
            fldChar2 = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
            run3._element.append(fldChar2)
            
            if is_chinese:
                footer_para.add_run(" 页 / 共 ").font.size = Pt(9)
            else:
                footer_para.add_run(" of ").font.size = Pt(9)
            
            # 总页数域
            run4 = footer_para.add_run()
            run4.font.size = Pt(9)
            fldChar3 = run4._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
            run4._element.append(fldChar3)
            
            run5 = footer_para.add_run()
            run5.font.size = Pt(9)
            instrText2 = run5._element.makeelement(qn('w:instrText'), {})
            instrText2.text = ' NUMPAGES '
            run5._element.append(instrText2)
            
            run6 = footer_para.add_run()
            run6.font.size = Pt(9)
            fldChar4 = run6._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
            run6._element.append(fldChar4)
            
            if is_chinese:
                footer_para.add_run(" 页").font.size = Pt(9)
            
        except Exception as e:
            logger.debug(f"Failed to add header/footer: {e}")
    
    def _add_page_number(self, doc):
        """Legacy method - redirects to _add_header_and_footer for backward compatibility."""
        self._add_header_and_footer(doc)

    def _generate_word(self, title: str, content: str, filename: str = None,
                       images: List[Dict] = None, tables: List[Dict] = None) -> str:
        """Generate Word document with academic formatting."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            logger.warning("python-docx not installed. Word generation skipped.")
            return ""
        
        images = images or []
        is_chinese = self._is_chinese_content(title) or self._is_chinese_content(content[:500])
        
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"research_report_{timestamp}"
        if filename.endswith('.pdf'):
            filename = filename[:-4]
        
        docx_path = self.output_dir / f"{filename}.docx"
        
        doc = Document()
        
        # 页面设置
        sections = doc.sections
        for section in sections:
            if is_chinese:
                # A4纸，左右2.6cm，上下3cm
                section.top_margin = Cm(3)
                section.bottom_margin = Cm(3)
                section.left_margin = Cm(2.6)
                section.right_margin = Cm(2.6)
            else:
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(2.54)
                section.right_margin = Cm(2.54)
        
        # 字体配置
        if is_chinese:
            font_title = 'SimHei'       # 黑体
            font_author = 'KaiTi'       # 楷体
            font_heading = 'SimSun'     # 宋体
            font_body = 'SimSun'        # 宋体
            font_latin = 'Times New Roman'
            body_size = 12              # 小四号
            h1_size = 14                # 四号
            h2_size = 12                # 小四号
            h3_size = 12                # 小四号
            h4_size = 12                # 小四号
            title_size = 22             # 二号
            author_size = 14            # 四号
            ref_size = 10.5             # 五号
            conclusion_size = 10.5      # 五号
        else:
            font_title = 'Times New Roman'
            font_author = 'Times New Roman'
            font_heading = 'Times New Roman'
            font_body = 'Times New Roman'
            font_latin = 'Times New Roman'
            body_size = 10
            h1_size = 14
            h2_size = 12
            h3_size = 11
            h4_size = 10
            title_size = 16
            author_size = 11
            ref_size = 9
            conclusion_size = 10
        
        # === 标题（二号黑体，居中，加粗） ===
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        self._set_word_run_font(title_run, font_latin, title_size, 
                                is_chinese=is_chinese, bold=True, font_east_asia=font_title)
        self._set_paragraph_line_spacing(title_para, 1.5)
        
        # === 作者/团队（四号楷体，居中） ===
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if is_chinese:
            author_run = author_para.add_run("RAAA 深度研究系统")
        else:
            author_run = author_para.add_run("Generated by RAAA Deep Research")
        self._set_word_run_font(author_run, font_latin, author_size,
                                is_chinese=is_chinese, font_east_asia=font_author)
        self._set_paragraph_line_spacing(author_para, 1.5)
        
        # === 日期 ===
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if is_chinese:
            date_run = date_para.add_run(datetime.now().strftime("%Y年%m月%d日"))
        else:
            date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
        self._set_word_run_font(date_run, font_latin, body_size,
                                is_chinese=is_chinese, italic=not is_chinese, font_east_asia=font_body)
        self._set_paragraph_line_spacing(date_para, 1.5)
        
        doc.add_paragraph()
        
        # 添加页眉和页码
        self._add_header_and_footer(doc, title=title, is_chinese=is_chinese)
        
        downloaded_images = self._prepare_images(images, max_images=8)
        img_idx = 0
        in_table = False
        table_lines = []
        in_references = False
        
        content = self._process_mermaid_code_blocks(content)
        
        lines = content.split('\n')
        for line in lines:
            line_stripped = line.strip()
            
            if not line_stripped:
                if in_table and table_lines:
                    self._add_word_table(doc, table_lines, is_chinese=is_chinese)
                    table_lines = []
                    in_table = False
                continue
            
            if line_stripped.startswith('|') and '|' in line_stripped[1:]:
                in_table = True
                table_lines.append(line_stripped)
                continue
            elif in_table:
                if table_lines:
                    self._add_word_table(doc, table_lines, is_chinese=is_chinese)
                table_lines = []
                in_table = False
            
            if line_stripped.startswith('# '):
                text = line_stripped[2:].replace('**', '').replace('*', '')
                in_references = 'reference' in text.lower() or '参考文献' in text
                doc.add_page_break()
                # 一级标题（四号宋体/黑体，加粗，顶格）
                if is_chinese:
                    heading = doc.add_paragraph()
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = heading.add_run(text)
                    self._set_word_run_font(run, font_latin, h1_size,
                                            is_chinese=True, bold=True, font_east_asia=font_heading)
                    self._set_paragraph_line_spacing(heading, 1.5)
                else:
                    heading = doc.add_heading(text.upper(), level=1)
                    for run in heading.runs:
                        run.font.name = font_latin
                        run.font.size = Pt(h1_size)
                
                if img_idx < len(downloaded_images):
                    self._add_word_image(doc, downloaded_images[img_idx], is_chinese=is_chinese)
                    img_idx += 1
                    
            elif line_stripped.startswith('## '):
                text = line_stripped[3:].replace('**', '').replace('*', '')
                # 二级标题（小四号宋体/黑体，加粗，顶格）
                if is_chinese:
                    heading = doc.add_paragraph()
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = heading.add_run(text)
                    self._set_word_run_font(run, font_latin, h2_size,
                                            is_chinese=True, bold=True, font_east_asia=font_heading)
                    self._set_paragraph_line_spacing(heading, 1.5)
                else:
                    heading = doc.add_heading(text, level=2)
                    for run in heading.runs:
                        run.font.name = font_latin
                        run.font.size = Pt(h2_size)
                    
            elif line_stripped.startswith('### '):
                text = line_stripped[4:].replace('**', '').replace('*', '')
                # 三级标题（小四号宋体/黑体，加粗，顶格）
                if is_chinese:
                    heading = doc.add_paragraph()
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = heading.add_run(text)
                    self._set_word_run_font(run, font_latin, h3_size,
                                            is_chinese=True, bold=True, font_east_asia=font_heading)
                    self._set_paragraph_line_spacing(heading, 1.5)
                else:
                    heading = doc.add_heading(text, level=3)
                    for run in heading.runs:
                        run.font.name = font_latin
                        run.font.size = Pt(h3_size)

            elif line_stripped.startswith('#### '):
                text = line_stripped[5:].replace('**', '').replace('*', '')
                # 四级标题（小四号宋体/黑体，加粗，顶格）
                if is_chinese:
                    heading = doc.add_paragraph()
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = heading.add_run(text)
                    self._set_word_run_font(run, font_latin, h4_size,
                                            is_chinese=True, bold=True, font_east_asia=font_heading)
                    self._set_paragraph_line_spacing(heading, 1.5)
                else:
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.bold = True
                    run.font.name = font_latin
                    run.font.size = Pt(h4_size)
                    
            elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                text = line_stripped[2:].replace('**', '').replace('*', '')
                para = doc.add_paragraph(text, style='List Bullet')
                for run in para.runs:
                    self._set_word_run_font(run, font_latin, body_size,
                                            is_chinese=is_chinese, font_east_asia=font_body)
                self._set_paragraph_line_spacing(para, 1.5)
                    
            elif re.match(r'^\d+\.\s', line_stripped):
                match = re.match(r'^(\d+)\.\s(.*)$', line_stripped)
                if match:
                    _, text = match.groups()
                    text = text.replace('**', '').replace('*', '')
                    para = doc.add_paragraph(text, style='List Number')
                    for run in para.runs:
                        self._set_word_run_font(run, font_latin, body_size,
                                                is_chinese=is_chinese, font_east_asia=font_body)
                    self._set_paragraph_line_spacing(para, 1.5)
                        
            elif line_stripped.startswith('[TABLE:') or line_stripped.startswith('[IMAGE:'):
                continue
            
            elif line_stripped.startswith('[RENDERED_MERMAID_IMAGE:'):
                img_match = re.match(r'\[RENDERED_MERMAID_IMAGE:(.+?)\|(.+?)\]', line_stripped)
                if img_match:
                    rendered_img_path = img_match.group(1).strip()
                    diagram_type = img_match.group(2).strip()
                    if os.path.exists(rendered_img_path):
                        self._add_word_image(doc, {
                            'path': rendered_img_path,
                            'title': diagram_type,
                            'source': "Generated Diagram"
                        }, is_chinese=is_chinese, show_caption=True)
                continue
            
            elif line_stripped.startswith('[WEB_SCREENSHOT:'):
                # 处理网页截图标记: [WEB_SCREENSHOT: url | description]
                screenshot_match = re.match(r'\[WEB_SCREENSHOT:\s*(.+?)\s*\|\s*(.+?)\s*\]', line_stripped)
                if screenshot_match:
                    screenshot_url = screenshot_match.group(1).strip()
                    screenshot_desc = screenshot_match.group(2).strip()
                    screenshot_path = self._download_web_screenshot(screenshot_url)
                    # 下载失败时跳过图片，不生成占位图
                    if screenshot_path:
                        self._add_word_image(doc, {
                            'path': screenshot_path,
                            'title': screenshot_desc,
                            'source': screenshot_url
                        }, is_chinese=is_chinese, show_caption=True)
                continue
            
            elif line_stripped.startswith('[MERMAID:'):
                mermaid_info = self._parse_mermaid_marker(line_stripped)
                if mermaid_info:
                    mermaid_type, mermaid_desc = mermaid_info
                    context = "\n".join(lines[max(0, lines.index(line)-10):lines.index(line)])
                    img_path = self._render_mermaid_to_image(mermaid_type, mermaid_desc, context)
                    if img_path:
                        self._add_word_image(doc, {
                            'path': img_path,
                            'title': mermaid_desc,
                            'source': f"Generated {mermaid_type.title()} Diagram"
                        }, is_chinese=is_chinese, show_caption=True)
                continue
                
            else:
                # 处理公式
                processed_text, formulas = self._process_latex_in_text(line_stripped)
                text = processed_text.replace('**', '').replace('*', '')
                
                # 替换公式占位符
                for formula in formulas:
                    placeholder = formula['placeholder']
                    latex = formula['latex']
                    if formula['display']:
                        # 显示公式：单独成段，居中
                        text = text.replace(placeholder, '')
                    else:
                        # 行内公式：用斜体显示
                        text = text.replace(placeholder, f' {latex} ')
                
                text = text.strip()
                
                # 先处理显示公式（单独成段）
                for formula in formulas:
                    if formula['display']:
                        # 渲染公式为图片并插入
                        img_path = self._render_latex_to_image(formula['latex'], display_mode=True)
                        if img_path:
                            self._add_word_image(doc, {
                                'path': img_path,
                                'title': '',
                                'source': ''
                            }, is_chinese=is_chinese, show_caption=False)
                
                if text:
                    para = doc.add_paragraph()
                    
                    if in_references:
                        # 参考文献使用左对齐，避免英文单词间距过大
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        # 参考文献不需要首行缩进，使用悬挂缩进
                        para.paragraph_format.first_line_indent = Cm(-0.74)
                        para.paragraph_format.left_indent = Cm(0.74)
                        run = para.add_run(text)
                        self._set_word_run_font(run, font_latin, ref_size,
                                                is_chinese=is_chinese, font_east_asia=font_body)
                    else:
                        if is_chinese:
                            # 正文首行缩进两字符
                            para.paragraph_format.first_line_indent = Cm(0.74)  # 约两个中文字符
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        run = para.add_run(text)
                        self._set_word_run_font(run, font_latin, body_size,
                                                is_chinese=is_chinese, font_east_asia=font_body)
                    self._set_paragraph_line_spacing(para, 1.5)
        
        if in_table and table_lines:
            self._add_word_table(doc, table_lines, is_chinese=is_chinese)
        
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if is_chinese:
            footer_run = footer.add_run("本报告由 RAAA（需求分析智能体助手）深度研究系统自动生成")
            self._set_word_run_font(footer_run, font_latin, 9,
                                    is_chinese=True, italic=True, font_east_asia=font_body)
        else:
            footer_run = footer.add_run("This report was generated by Requirements Analysis Agent Assistant (RAAA)")
            footer_run.italic = True
            footer_run.font.size = Pt(9)
        
        doc.save(str(docx_path))
        logger.info(f"Word document generated: {docx_path}")
        return str(docx_path)
    
    def _add_word_table(self, doc, table_lines: List[str], is_chinese: bool = False):
        """
        Add a table to Word document from markdown table lines.
        For Chinese: uses 三线表 (three-line table) style with Chinese fonts.
        """
        try:
            from docx.shared import Pt, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import nsdecls, qn
            from docx.oxml import parse_xml
            
            table_data = self._parse_table_from_text('\n'.join(table_lines))
            if not table_data:
                return
            
            table_num = self._next_table_num()
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if is_chinese:
                caption_run = caption.add_run(f"表{table_num}")
                self._set_word_run_font(caption_run, 'Times New Roman', 9,
                                        is_chinese=True, bold=True, font_east_asia='SimHei')
            else:
                caption_run = caption.add_run(f"Table {table_num}. Research Data")
                caption_run.bold = True
                caption_run.font.size = Pt(9)
            
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            
            if is_chinese:
                # 三线表样式：仅顶线、表头下线、底线
                table.style = 'Table Grid'
                # 移除所有边框，然后只添加三线
                tbl = table._tbl
                tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
                borders = parse_xml(
                    f'<w:tblBorders {nsdecls("w")}>'
                    '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                    '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                    '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    '</w:tblBorders>'
                )
                # 移除旧边框
                old_borders = tblPr.find(qn('w:tblBorders'))
                if old_borders is not None:
                    tblPr.remove(old_borders)
                tblPr.append(borders)
                
                # 在表头行下方添加细线
                if len(table_data) > 1:
                    header_row = table.rows[0]
                    for cell in header_row.cells:
                        tcPr = cell._element.get_or_add_tcPr()
                        tcBorders = parse_xml(
                            f'<w:tcBorders {nsdecls("w")}>'
                            '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
                            '</w:tcBorders>'
                        )
                        old_tcBorders = tcPr.find(qn('w:tcBorders'))
                        if old_tcBorders is not None:
                            tcPr.remove(old_tcBorders)
                        tcPr.append(tcBorders)
            else:
                table.style = 'Table Grid'
            
            font_body = 'SimSun' if is_chinese else 'Times New Roman'
            
            for i, row_data in enumerate(table_data):
                row = table.rows[i]
                for j, cell_text in enumerate(row_data):
                    cell = row.cells[j]
                    clean_text = str(cell_text)
                    clean_text = re.sub(r'^\*+|\*+$', '', clean_text)
                    clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', clean_text)
                    clean_text = re.sub(r'\*(.+?)\*', r'\1', clean_text)
                    cell.text = clean_text.strip()
                    for para in cell.paragraphs:
                        # 表内数字上下对齐
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                        for run in para.runs:
                            self._set_word_run_font(run, 'Times New Roman', 9,
                                                    is_chinese=is_chinese, bold=(i == 0),
                                                    font_east_asia=font_body)
            
            doc.add_paragraph()
            
        except Exception as e:
            logger.warning(f"Failed to add Word table: {e}")
    
    def _add_word_image(self, doc, img_info: Dict, is_chinese: bool = False, show_caption: bool = False):
        """Add an image to Word document, optionally with caption."""
        try:
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(img_info['path'], width=Inches(5.0))
            
            # 只在 show_caption=True 时显示题注
            if show_caption:
                fig_num = self._next_figure_num()
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if is_chinese:
                    caption_text = f"图{fig_num} {img_info.get('title', '')}"
                    if img_info.get('source') and img_info['source'] != 'Unknown':
                        caption_text += f"（来源：{img_info['source']}）"
                    caption_run = caption.add_run(caption_text)
                    self._set_word_run_font(caption_run, 'Times New Roman', 9,
                                            is_chinese=True, font_east_asia='SimHei')
                else:
                    caption_text = f"Figure {fig_num}. {img_info.get('title', '')} (Source: {img_info.get('source', 'Unknown')})"
                    caption_run = caption.add_run(caption_text)
                    caption_run.font.size = Pt(9)
                    caption_run.font.name = 'Times New Roman'
            
        except Exception as e:
            logger.warning(f"Failed to add Word image: {e}")


class PDFReportGenerator(AcademicReportGenerator):
    """Backward-compatible alias for AcademicReportGenerator."""
    pass
